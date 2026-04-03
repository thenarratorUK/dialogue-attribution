import streamlit as st
import re
import os
import json
import tempfile
import mammoth
import docx
import base64
import io
import csv
from pathlib import Path
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup, NavigableString
from collections import Counter
from streamlit_theme import st_theme
import html
from datetime import datetime, timezone


# -----------------------------
# PDF export (HTML -> PDF)
# -----------------------------
@st.cache_data(show_spinner=False)
def is_single_titlecase_speaker_label(label_text, next_text=""):
    """True for a single-word Titlecase speaker label followed by ':' and then whitespace.

    Handles DOCX run boundaries where the colon/space may be in a separate run with different styling.

    Matches:
      label_text='Pixel'      next_text=': '   -> True
      label_text='Pixel:'     next_text=' '    -> True
      label_text='Pixel: '    next_text=''     -> True
      label_text='Friedrich'  next_text=':\t' -> True

    Does NOT match:
      '12:34'
      'pixel'
      'PIXEL'
    """
    if label_text is None:
        return False

    s = str(label_text)

    def norm_ws(x: str) -> str:
        return (x or "").replace("\u00A0", " ").replace("\u202F", " ").replace("\u2009", " ").replace("\u200A", " ").replace("\u200B", "")

    s = norm_ws(s)
    nxt = norm_ws("" if next_text is None else str(next_text))

    has_colon_and_ws = (
        s.endswith(": ") or
        s.endswith(":\t") or
        (s.endswith(":") and (nxt[:1].isspace() or nxt.startswith(" ")))
    )

    has_name_then_colon_ws_in_next = (
        nxt.startswith(":") and len(nxt) >= 2 and nxt[1].isspace()
    )

    if has_colon_and_ws:
        core = s.rstrip()
        if not core.endswith(":"):
            return False
        name = core[:-1]
    elif has_name_then_colon_ws_in_next:
        name = s.rstrip()
    else:
        return False

    if not re.fullmatch(r"[A-Z][a-z]+", name or ""):
        return False

    return True


def render_html_to_pdf_bytes(html_str: str, base_url: str) -> bytes:
    """Render an HTML string to a PDF (bytes).

    This uses WeasyPrint if available. The extra CSS forces print colour fidelity so
    background colours are preserved in the resulting PDF.
    """
    from weasyprint import HTML, CSS  # type: ignore
    extra_css = CSS(string="""
        * { print-color-adjust: exact; }
        @page { size: A4; margin: 18mm; }
    """)
    return HTML(string=html_str, base_url=base_url).write_pdf(stylesheets=[extra_css])

def encode_font_base64(path: str) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def build_font_face_css(fontsel: str, embed_base64: bool = False) -> str:
    """
    Return one or more @font-face CSS blocks for the selected font, or "" if not needed.
    If embed_base64 is True, embed the font files as Base64 data URLs.

    Supported families:
      - Lexend: Lexend-VariableFont_wght.ttf (variable weight, no italics)
      - Gentium Basic: Regular / Italic / Bold / Bold-Italic (TTF)
      - OpenDyslexic: Regular / Italic / Bold / Bold-Italic (OTF)
    Other fonts (Avenir, Helvetica, etc.) are treated as system fonts only.
    """

    def _face(
        family: str,
        path: str,
        weight: str,
        style: str,
        fmt: str,
        mime_subtype: str,
    ) -> str:
        """
        Build a single @font-face rule for one file.
        fmt: 'truetype' or 'opentype'
        mime_subtype: 'ttf' or 'otf'
        """
        if embed_base64:
            try:
                b64 = encode_font_base64(path)
            except FileNotFoundError:
                return ""
            src = f"data:font/{mime_subtype};base64,{b64}"
        else:
            src = path

        return f"""
@font-face {{
  font-family: '{family}';
  src: url('{src}') format('{fmt}');
  font-weight: {weight};
  font-style: {style};
}}
"""

    rules: list[str] = []

    # ---------------- Lexend (variable font, TTF) ----------------
    if fontsel == "Lexend":
        # Single variable font covering weights 100–900, normal style only.
        # Italic will be synthetic since there is no italic file.
        rules.append(
            _face(
                family="Lexend",
                path="fonts/Lexend-VariableFont_wght.ttf",
                weight="100 900",       # variable font weight range
                style="normal",
                fmt="truetype",
                mime_subtype="ttf",
            )
        )

    # --------------- Gentium Basic (full family, TTF) -------------
    elif fontsel == "Gentium Basic":
        # Regular
        rules.append(
            _face(
                family="Gentium Basic",
                path="fonts/GentiumBasic-Regular.ttf",
                weight="400",
                style="normal",
                fmt="truetype",
                mime_subtype="ttf",
            )
        )
        # Italic
        rules.append(
            _face(
                family="Gentium Basic",
                path="fonts/GentiumBasic-Italic.ttf",
                weight="400",
                style="italic",
                fmt="truetype",
                mime_subtype="ttf",
            )
        )
        # Bold
        rules.append(
            _face(
                family="Gentium Basic",
                path="fonts/GentiumBasic-Bold.ttf",
                weight="700",
                style="normal",
                fmt="truetype",
                mime_subtype="ttf",
            )
        )
        # Bold-Italic
        rules.append(
            _face(
                family="Gentium Basic",
                path="fonts/GentiumBasic-Bold-Italic.ttf",
                weight="700",
                style="italic",
                fmt="truetype",
                mime_subtype="ttf",
            )
        )

    # --------------- OpenDyslexic (full family, OTF) --------------
    elif fontsel in ("OpenDyslexic", "Open Dyslexic"):
        # Regular
        rules.append(
            _face(
                family="OpenDyslexic",
                path="fonts/OpenDyslexic-Regular.otf",
                weight="400",
                style="normal",
                fmt="opentype",
                mime_subtype="otf",
            )
        )
        # Italic
        rules.append(
            _face(
                family="OpenDyslexic",
                path="fonts/OpenDyslexic-Italic.otf",
                weight="400",
                style="italic",
                fmt="opentype",
                mime_subtype="otf",
            )
        )
        # Bold
        rules.append(
            _face(
                family="OpenDyslexic",
                path="fonts/OpenDyslexic-Bold.otf",
                weight="700",
                style="normal",
                fmt="opentype",
                mime_subtype="otf",
            )
        )
        # Bold-Italic
        rules.append(
            _face(
                family="OpenDyslexic",
                path="fonts/OpenDyslexic-Bold-Italic.otf",
                weight="700",
                style="italic",
                fmt="opentype",
                mime_subtype="otf",
            )
        )

    # --------------- System fonts (no embedding) -------------------
    else:
        # Avenir, Helvetica, Arial, Georgia, Times New Roman, Courier New, etc.
        # Use system-installed fonts only; no @font-face needed.
        return ""

    return "".join(rules)


def normalize_font_family(fontsel: str) -> str:
    """Normalize UI labels / legacy values to CSS font-family names."""
    return "OpenDyslexic" if fontsel == "Open Dyslexic" else fontsel


def font_label_to_css_family(font_label: str) -> str:
    """Map UI dropdown labels to the CSS font-family values used by the app."""
    if font_label == "Open Dyslexic":
        return "OpenDyslexic"
    return font_label


def css_family_to_font_label(css_family: str) -> str:
    """Map stored CSS font-family values back to UI dropdown labels."""
    if css_family in ("Open Dyslexic", "OpenDyslexic"):
        return "Open Dyslexic"
    return css_family

_MOJIBAKE_FIXES = {
    # common utf8->latin1
    "â€˜": "‘", "â€™": "’", "â€œ": "“", "â€": "”",
    "â€“": "–", "â€”": "—", "â€¦": "…",
    # mac/other weird
    "‚Äò": "‘", "‚Äô": "’", "‚Äú": "“", "‚Äù": "”",
    "‚Äî": "—", "‚Äì": "–", "‚Ä¶": "…",
    "Ä¶": "…",
    # stray nbsp-ish
    "Â": "",
}

def render_brand_header(logo_width_px: int = 200):
    """Render the brand header (logo left, text right). Uses logo_alt.png when Streamlit theme is dark."""
    left, middle, right = st.columns([1, 1, 1], vertical_alignment="center")

    with left:
        logo_path = Path(__file__).with_name("logo.png")
        logo_alt_path = Path(__file__).with_name("logo_alt.png")

        if logo_path.exists():
            theme = st_theme() or {}
            base = (theme.get("base") or "").lower()

            # If Streamlit is in dark mode and alt logo exists, use it; otherwise use default.
            chosen = logo_alt_path if (base == "dark" and logo_alt_path.exists()) else logo_path
            st.image(str(chosen), width=logo_width_px)

    with right:
        st.markdown(
            'Created by David Winter  \n("The Narrator")  \nhttps://www.thenarrator.co.uk  \nReadme: [Link](https://github.com/thenarratorUK/dialogue-attribution/blob/main/Readme.md)'
        )

    st.markdown("---")
    
def render_font_preview(font_options: list[str]):
    """Show each font name rendered in its own font as a quick preview."""
    preview_lines = []
    for font_name in font_options:
        css_font = normalize_font_family("OpenDyslexic" if font_name == "Open Dyslexic" else font_name)
        preview_lines.append(
            f'<div style="font-family: {css_font}, sans-serif; margin: 0.1rem 0;">{html.escape(font_name)}</div>'
        )

    st.markdown("**Font preview:**", help="Streamlit selectbox options cannot be styled per row, so previews are shown below.")
    st.markdown("\n".join(preview_lines), unsafe_allow_html=True)

def _fix_mojibake(s: str) -> str:
    for k, v in _MOJIBAKE_FIXES.items():
        s = s.replace(k, v)
    return s
    
def build_csv_from_docx_json_and_quotes():
    """
    Rebuild the paragraph JSON from the DOCX and use quotes_lines
    to generate a CSV with Speaker + Line + FileName using the
    search/trim loop.

    Changes vs previous version:
      - Match on HTML-stripped, mojibake-fixed paragraph text.
      - Skip 'Do Not Read:' lines from quotes.txt.
      - Keep a minimal fallback but avoid duplicate narration
        for unmatched quote segments.
    """
    import os, json, re, io, csv

    # If we're in Script mode, build the CSV directly from quotes.txt
    if st.session_state.get("content_type", "Book") == "Script":
        quotes_lines = st.session_state.get("quotes_lines") or []
        canonical_map = st.session_state.get("canonical_map") or {}

        # Parse each quotes.txt line: "123. Speaker: Dialogue"
        pattern = re.compile(r"^\s*([0-9]+(?:[a-zA-Z]+)?)\.\s+([^:]+):\s*(?:[“\"])?(.+?)(?:[”\"])?\s*$")

        rows: list[tuple[str, str]] = []

        for raw_line in quotes_lines:
            line = raw_line.strip()
            if not line:
                continue
            m = pattern.match(line)
            if not m:
                continue
            _, speaker_raw, quote = m.groups()
            effective = smart_title(speaker_raw)
            norm = normalize_speaker_name(effective)
            canonical = canonical_map.get(norm, effective)
            text_part = quote.strip()
            if not text_part:
                continue
            rows.append((canonical, text_part))

        # Build CSV: same filename pattern as the book workflow
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(["Speaker", "Line", "FileName"])

        def normalise_speaker_name_local(s: str) -> str:
            if s and s.strip().lower() == "error":
                return "Narration"
            return s

        for idx, (speaker, line) in enumerate(rows, start=1):
            speaker_clean = normalise_speaker_name_local(speaker)
            line_clean = normalize_text(line)
            num = f"{idx:05d}"
            safe_speaker = re.sub(r"\s+", "", speaker_clean) or "Narration"
            filename = f"{num}_{safe_speaker}_TakeX"
            writer.writerow([speaker_clean, line_clean, filename])

        return buf.getvalue().encode("utf-8")

    # Ensure the JSON is up to date for the current DOCX
    write_paragraph_json_for_session()
    json_path = st.session_state.get("d_json_path")
    if not json_path or not os.path.exists(json_path):
        return b""

    # --- load and clean paragraph list ---
    with open(json_path, "r", encoding="utf-8") as f:
        raw_paragraphs = json.load(f)

    def strip_tags(text: str) -> str:
        # Remove simple HTML-like tags but do NOT normalise whitespace here
        return re.sub(r"<[^>]*>", "", text or "")

    # Plain, mojibake-fixed paragraphs used for all matching and output
    paragraphs_plain = [_fix_mojibake(strip_tags(p)) for p in raw_paragraphs]

    quotes_lines = st.session_state.get("quotes_lines") or []

    def normalise_speaker_name(s: str) -> str:
        # We still honour Error->Narration, but do it at the very end as well
        if s and s.strip().lower() == "error":
            return "Narration"
        return s

    rows: list[tuple[str, str]] = []      # (Speaker, Line)
    remaining = list(paragraphs_plain)    # trimmed as we go (plain text only)
    unmatched_segments: list[str] = []    # track text we emitted via fallback

    # =============== CORE LOOP THROUGH QUOTES.TXT ===============
    for raw_line in quotes_lines:
        line = raw_line.strip()
        if not line:
            continue

        # Remove leading "2621. " style numbering
        line_wo_num = re.sub(r"^\s*\d+\.\s*", "", line)

        if ":" not in line_wo_num:
            continue

        speaker_part, quote_part = line_wo_num.split(":", 1)
        speaker_raw = speaker_part.strip()
        quote_text = quote_part.strip()
        if not quote_text:
            continue

        # Skip "Do Not Read:" lines entirely
        if speaker_raw.lower().startswith("do not read"):
            continue

        speaker_norm = normalise_speaker_name(speaker_raw)

        # Use mojibake-fixed quote text for matching
        quote_match = _fix_mojibake(quote_text)

        # ---- FIND QUOTE IN REMAINING PARAGRAPHS (PLAIN TEXT) ----
        found_idx = -1
        found_pos = -1

        for idx, para in enumerate(remaining):
            pos = para.find(quote_match)
            if pos != -1:
                found_idx = idx
                found_pos = pos
                break

        if found_idx == -1:
            # Fallback: emit the quote from TXT, but remember it so we don't
            # also emit the same text as Narration later.
            unmatched_segments.append(quote_match)
            rows.append((speaker_norm, quote_match))
            continue

        # ---- 1. PARAGRAPHS BEFORE MATCH = NARRATION ----
        before_paras = remaining[:found_idx]
        for pre in before_paras:
            plain = pre.strip()
            if plain and plain not in unmatched_segments:
                rows.append(("Narration", plain))

        # ---- 2. TEXT BEFORE QUOTE IN MATCHING PARAGRAPH ----
        current = remaining[found_idx]
        before = current[:found_pos]
        before_plain = before.strip()
        if before_plain and before_plain not in unmatched_segments:
            rows.append(("Narration", before_plain))

        # ---- 3. THE QUOTE ITSELF ----
        rows.append((speaker_norm, quote_match))

        # ---- 4. TRIM REMAINING TO AFTER QUOTE ----
        after = current[found_pos + len(quote_match):]
        new_remaining: list[str] = []
        if after:
            new_remaining.append(after)
        new_remaining.extend(remaining[found_idx + 1:])
        remaining = new_remaining

    # ---- FINAL TAIL NARRATION ----
    for tail in remaining:
        plain = tail.strip()
        if plain and plain not in unmatched_segments:
            rows.append(("Narration", plain))

    # =============== BUILD CSV WITH FILENAME COLUMN ===============
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Speaker", "Line", "FileName"])

    for idx, (speaker, line) in enumerate(rows, start=1):
        # Mojibake cleanup at the very end
        speaker_clean = _fix_mojibake(speaker)
        line_clean = _fix_mojibake(line)

        # Error -> Narration normalisation
        if speaker_clean.strip().lower() == "error":
            speaker_clean = "Narration"

        num = f"{idx:05d}"
        safe_speaker = re.sub(r"\s+", "", speaker_clean) or "Narration"
        filename = f"{num}_{safe_speaker}_TakeX"

        writer.writerow([speaker_clean, line_clean, filename])

    return buf.getvalue().encode("utf-8")

def trim_paragraph_cache_before_previous(previous_html: str):
    # Intentionally disabled in test flow:
    # destructive cache trimming is unsafe when duplicate paragraphs exist,
    # because content-equality matching cannot identify the exact prior location.
    return False

def neutralize_markdown_in_html(html_s: str) -> str:
    try:
        soup = BeautifulSoup(html_s, "html.parser")
        for t in list(soup.find_all(string=True)):
            # Skip script/style tags
            if getattr(t.parent, "name", "").lower() in ("script", "style"):
                continue
            # Replace literal '*' and '_' with HTML entities so st.markdown doesn't italicise them.
            new_txt = t.replace("*", r"\*").replace("_", r"\_")
            if new_txt != t:
                t.replace_with(new_txt)
        return str(soup)
    except Exception:
        # Fallback: raw replacements (may affect tags if present, but better than italics)
        return html_s.replace("*", "&#42;").replace("_", "&#95;")

# Import components for HTML embedding.
import streamlit.components.v1 as components

def build_d_paragraphs_html(docx_path):
    import docx
    import html
    try:
        doc = docx.Document(docx_path)
    except Exception:
        return []
    def wrap(txt, b, i, u):
        s = html.escape(txt or "")
        if not s:
            return s
        if u: s = f"<u>{s}</u>"
        if i: s = f"<i>{s}</i>"
        if b: s = f"<b>{s}</b>"
        return s
    out = []
    # ensure `re` is imported at top of file: import re
    for p in doc.paragraphs:
        if not p.runs:
            candidate = html.escape(p.text or "")
        else:
            candidate = "".join(
                wrap(r.text, r.bold, r.italic, r.underline) for r in p.runs
            )
    
        # Drop empty/whitespace-only paragraphs (ignoring any HTML tags)
        plain = re.sub(r"<[^>]*>", "", candidate)  # strip tags for the emptiness check
        if not re.search(r"\S", plain):            # no non-whitespace char
            continue
    
        out.append(candidate)
    return out



def get_context_for_dialogue_json_only(dialogue: str, occurrence_target: int = 1, start_paragraph_index: int = 0):
    try:
        djson_path = st.session_state.get('d_json_path')
        if not djson_path or not os.path.exists(djson_path):
            return None
        with open(djson_path, "r", encoding="utf-8") as f:
            paragraphs_html = json.load(f)  # list[str]
    except Exception:
        return None

    def soup_text(html_s: str) -> str:
        try:
            soup = BeautifulSoup(html_s, "html.parser")
            return soup.get_text() or ""
        except Exception:
            return html_s

    dlg = dialogue
    m_q = re.search(r'[“"]([^”"]+)[”"]', dlg) or re.search(r"[‘']([^’']+)[’']", dlg)
    dialogue_to_highlight = m_q.group(1) if m_q else dlg
    normalized_highlight = normalize_text(dialogue_to_highlight).lower()

    if not normalized_highlight.strip():
        occurrence_target = 1

    plain_paras = [soup_text(p) for p in paragraphs_html]

    try:
        start_idx = max(0, int(start_paragraph_index or 0))
    except Exception:
        start_idx = 0
    if start_idx >= len(plain_paras):
        start_idx = 0

    cumulative = 0
    chosen_idx = None
    within_para_target = 1

    for idx in range(start_idx, len(plain_paras)):
        para_plain = plain_paras[idx]
        para_norm = normalize_text(para_plain).lower()
        count_here = count_with_boundaries_ci(para_norm, normalized_highlight) if normalized_highlight else 0
        if count_here > 0:
            if cumulative + count_here >= occurrence_target:
                chosen_idx = idx
                within_para_target = occurrence_target - cumulative
                break
            cumulative += count_here

    if chosen_idx is None:
        for idx in range(start_idx, len(plain_paras)):
            para_plain = plain_paras[idx]
            if find_with_boundaries_ci(normalize_text(para_plain).lower(), normalized_highlight, 0) is not None:
                chosen_idx = idx
                within_para_target = 1
                break

    if chosen_idx is None:
        return None

    ctx = {}
    ctx["paragraph_index"] = chosen_idx
    if chosen_idx > 0:
        ctx["previous"] = paragraphs_html[chosen_idx - 1]

    try:
        soup = BeautifulSoup(paragraphs_html[chosen_idx], "html.parser")
                # --- begin: tag-stripped match + bold in original HTML (early-return if applied) ---
        # 1) Build a tag-stripped view for matching plus normalized index-map
        plain_text = soup.get_text()
        plain_norm, norm_to_raw = normalize_text_with_index_map(plain_text)

        # 2) Find the m-th (within_para_target) occurrence on normalized plain text
        _span_norm = nth_span_with_boundaries_ci(plain_norm.lower(), normalized_highlight, within_para_target)
        _span = None
        if _span_norm is not None and norm_to_raw:
            ns, ne = _span_norm
            if 0 <= ns < len(norm_to_raw) and 0 <= (ne - 1) < len(norm_to_raw):
                raw_start = norm_to_raw[ns]
                raw_end = norm_to_raw[ne - 1] + 1
                _span = (raw_start, raw_end)
        
        # 3) If found, map the [start:end) range back onto the original soup's text nodes and wrap with <b>
        if _span is not None:
            _start, _end = _span
            running = 0
            # Walk text nodes in order; split and wrap overlap segments with <b>
            for tnode in list(soup.find_all(string=True)):
                p = getattr(tnode.parent, "name", "").lower()
                if p in ("script", "style"):
                    continue
                text = str(tnode)
                length = len(text)
                node_start = running
                node_end = running + length
        
                # Does this node overlap the target [start, end)?
                if node_end > _start and node_start < _end:
                    # overlap in this node
                    ov_start = max(_start, node_start)
                    ov_end   = min(_end, node_end)
                    rel_start = ov_start - node_start
                    rel_end   = ov_end   - node_start
        
                    before = text[:rel_start]
                    match  = text[rel_start:rel_end]
                    after  = text[rel_end:]
        
                    from bs4 import NavigableString
                    new_nodes = []
                    if before:
                        new_nodes.append(NavigableString(before))
                    btag = soup.new_tag("b")
                    btag.string = match
                    new_nodes.append(btag)
                    if after:
                        new_nodes.append(NavigableString(after))
        
                    # Replace this text node with the split nodes
                    tnode.replace_with(*new_nodes)
        
                    # IMPORTANT: if match spans multiple consecutive text nodes,
                    # decrease the remaining range and continue across nodes
                    # Update the remaining global range for subsequent nodes:
                    # we wrapped [ov_start, ov_end), so advance _start to ov_end
                    _start = ov_end
                    if _start >= _end:
                        break  # fully wrapped
        
                    # Adjust running to reflect we've consumed this node
                    running = node_end
                    continue
        
                running += length
        
            ctx["current"] = str(soup)
            if chosen_idx + 1 < len(paragraphs_html):
                ctx["next"] = paragraphs_html[chosen_idx + 1]
            return ctx
        # --- end: tag-stripped match + bold in original HTML ---
        pattern = re.compile(re.escape(dialogue_to_highlight), re.IGNORECASE)

        global_counter = 0
        replaced_flag = False

        def replace_only_mth(node):
            nonlocal global_counter, replaced_flag
            text_val = str(node)
            parts = []
            last = 0
            for m in pattern.finditer(text_val):
                global_counter += 1
                parts.append(text_val[last:m.start()])
                seg = m.group(0)
                if (not replaced_flag) and global_counter == within_para_target:
                    b = soup.new_tag("b")
                    b.string = seg
                    parts.append(b)
                    replaced_flag = True
                else:
                    parts.append(seg)
                last = m.end()
            parts.append(text_val[last:])
            if len(parts) > 1:
                anchor = None
                for i, part in enumerate(parts):
                    new_node = soup.new_string(part) if isinstance(part, str) else part
                    if i == 0:
                        node.replace_with(new_node)
                        anchor = new_node
                    else:
                        anchor.insert_after(new_node)
                        anchor = new_node

        for tnode in list(soup.find_all(string=True)):
            parent_name = getattr(tnode.parent, "name", "").lower()
            if parent_name in ("script", "style"):
                continue
            if replaced_flag:
                continue
            replace_only_mth(tnode)

        if not replaced_flag:
            global_counter = 0
            replaced_flag = False
            for tnode in list(soup.find_all(string=True)):
                parent_name = getattr(tnode.parent, "name", "").lower()
                if parent_name in ("script", "style"):
                    continue
                replace_only_mth(tnode)
                if replaced_flag:
                    break

        ctx["current"] = str(soup)
    except Exception:
        ctx["current"] = paragraphs_html[chosen_idx]

    if chosen_idx + 1 < len(paragraphs_html):
        ctx["next"] = paragraphs_html[chosen_idx + 1]

    return ctx


#def compute_and_set_d_json_path():
#    """Ensure JSON paragraph cache exists and set st.session_state['d_json_path'].
#    JSON is written alongside the DOCX using the book name. No DOCX fallback in Step 2.
#    """
#    try:
#        docx_path = st.session_state.get("docx_path")
#        book_name = st.session_state.get("book_name") or "document"
#        if not docx_path or not os.path.exists(docx_path):
#            return
#       folder = os.path.dirname(docx_path)
#        json_path = os.path.join(folder, f"{book_name}.json")
#        paragraphs = build_d_paragraphs_html(docx_path)
#        with open(json_path, "w", encoding="utf-8") as f:
#            json.dump(paragraphs, f, ensure_ascii=False, indent=2)
#        st.session_state['d_json_path'] = json_path
#    except Exception:
#        # Do not fall back; if JSON generation fails, leave preview blank.
#        pass

def write_paragraph_json_for_session():
    """Create or overwrite [userkey]-[bookname].json in the working directory
    from st.session_state['docx_path'], set st.session_state['d_json_path'],
    and never fall back to DOCX later. This runs only when 'Start Processing' is clicked.
    """
    try:
        import os, json
        docx_path = st.session_state.get('docx_path')
        if not docx_path or not os.path.exists(docx_path):
            return
        userkey = st.session_state.get('userkey') or "User"
        book_name = st.session_state.get('book_name') or "document"
        json_path = os.path.join(os.getcwd(), f"{userkey}-{book_name}.json")
        paragraphs = build_d_paragraphs_html(docx_path)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(paragraphs, f, ensure_ascii=False, indent=2)
        st.session_state['d_json_path'] = json_path
    except Exception:
        # Do not silently recreate elsewhere; leave preview blank if this fails.
        pass

#def ensure_d_json(docx_path, quotes_path):
#    """Deprecated: use write_paragraph_json_for_session(). Keeping for backward compatibility."""
#    write_paragraph_json_for_session()

# Ensure a default font selection in session_state
if "fontsel" not in st.session_state:
    st.session_state.fontsel = "Avenir"

if "fontsel_label" not in st.session_state:
    st.session_state.fontsel_label = css_family_to_font_label(
        normalize_font_family(st.session_state.fontsel)
    )

fontsel = normalize_font_family(
    font_label_to_css_family(st.session_state.get("fontsel_label", "Avenir"))
)
st.session_state.fontsel = fontsel

# Apply the selected font globally across the full app, including start page.
# For custom bundled fonts, use Base64 so Streamlit can load fonts reliably.
font_face_css = build_font_face_css(fontsel, embed_base64=True)
# Inject custom CSS
custom_css = f"""
<style>
{font_face_css}
:root {{
  --primary-color: #008080;      /* Teal */
  --primary-hover: #007070;
  --background-color: #fdfdfd;
  --text-color: #222222;
  --card-background: #ffffff;
  --card-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
  --border-radius: 10px;
  --font-family: '{fontsel}', sans-serif;
  --accent-color: #ff9900;
}}

/* Global Styles */
body {{
  background-color: var(--background-color);
  font-family: var(--font-family);
  color: var(--text-color);
  margin: 0;
  padding: 0;
}}

/* Streamlit renders most UI inside .stApp, so set font there too. */
.stApp, .stApp p, .stApp span, .stApp label, .stApp li,
.stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp h5, .stApp h6,
.stApp div[data-baseweb="select"] *,
.stApp div[data-baseweb="input"] *,
.stApp div[data-baseweb="textarea"] *,
.stApp [data-testid="stMarkdownContainer"] * {{
  font-family: var(--font-family) !important;
}}

h1, h2, h3, h4, h5, h6 {{
  color: var(--text-color);
  font-weight: 700;
  margin-bottom: 0.5em;
}}

/* Button Styles */
div.stButton > button {{
  background-color: var(--primary-color);
  color: #ffffff;
  border: none;
  padding: 0.75em 1.25em;
  border-radius: var(--border-radius);
  cursor: pointer;
  transition: background-color 0.3s ease, transform 0.2s;
}}

div.stButton > button:hover {{
  background-color: var(--primary-hover);
  transform: translateY(-2px);
}}

/* Card/Container Styling */
.custom-container {{
  background: var(--card-background);
  padding: 2em;
  border-radius: var(--border-radius);
  box-shadow: var(--card-shadow);
  margin-bottom: 2em;
}}

.css-1d391kg {{
  background: var(--card-background);
  padding: 1em;
  border-radius: var(--border-radius);
  box-shadow: var(--card-shadow);
}}

/* Form Element Styling */
input, select, textarea {{
  border: 1px solid #ddd;
  border-radius: 4px;
  padding: 0.5em;
  font-size: 1em;
}}

input:focus, select:focus, textarea:focus {{
  outline: none;
  border-color: var(--primary-color);
  box-shadow: 0 0 5px rgba(0, 128, 128, 0.3);
}}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ---------------------------
# Global Constants & Helper Functions
# ---------------------------

COLOR_PALETTE = {
    "dark grey": (180, 178, 179, 1, "rgb(30, 28, 29)"),
    "burgundy": (207, 156, 153, 1, "rgb(55, 4, 1)"),
    "red": (255, 153, 153, 1, "rgb(107, 9, 6)"),
    "orange": (255, 198, 153, 1, "rgb(119, 62, 17)"),
    "yellow": (255, 241, 153, 1, "rgb(107, 93, 5)"),
    "dark yellow": (230, 223, 153, 1, "rgb(78, 71, 1)"),
    "brown": (205, 192, 173, 1, "rgb(61, 48, 29)"),
    "silver": (243, 243, 243, 1, "rgb(96, 96, 96)"),
    "light green": (160, 255, 153, 1, "rgb(20, 115, 13)"),
    "dark green": (157, 192, 152, 1, "rgb(7, 42, 2)"),
    "turquoise": (153, 255, 234, 1, "rgb(3, 105, 84)"),
    "light blue": (153, 248, 254, 1, "rgb(9, 104, 110)"),
    "bright blue": (152, 166, 255, 1, "rgb(4, 18, 107)"),
    "dark blue": (152, 166, 255, 1, "rgb(4, 18, 107)"),
    "navy blue": (153, 162, 201, 1, "rgb(4, 13, 52)"),
    "dark purple": (219, 173, 222, 1, "rgb(68, 22, 71)"),
    "light purple": (231, 203, 254, 1, "rgb(81, 53, 104)"),
    "bright pink": (255, 153, 236, 1, "rgb(112, 10, 93)"),
    "light pink": (254, 238, 248, 1, "rgb(103, 87, 97)"),
    "pale pink": (254, 238, 248, 1, "rgb(103, 87, 97)"),
    "wine": (234, 184, 185, 1, "rgb(90, 40, 41)"),
    "lime": (230, 244, 182, 1, "rgb(81, 95, 33)"),
    "none": (134, 8, 0, 1.0, "rgb(134, 8, 0)"),
    "do not read": (0, 0, 0, 1, "rgb(100, 100, 100)"),
    "error": (0, 0, 0, 0, "")  # For "Error": transparent background, no text color override.
}


# ==== STEP 0: Userkey Entry ====
if "userkey" not in st.session_state:
    st.session_state.userkey = ""

if "step" not in st.session_state:
    st.session_state.step = 0

if "content_type" not in st.session_state:
    st.session_state.content_type = "Book"
if "review_events" not in st.session_state:
    st.session_state.review_events = []

# ========= STEP 0: User Identification =========

if st.session_state.step == 0:
    render_brand_header()
    st.title("Welcome to Scripter")
    st.write(
        "Please enter a unique identifier. This can be any memorable username or passphrase. "
        "It must be unique to you—do not share it."
    )

    # Font selection for UI + exported HTML
    font_options = [
        "Avenir",
        "Helvetica",
        "Arial",
        "Georgia",
        "Times New Roman",
        "Courier New",
        "Open Dyslexic",
        "Gentium Basic",
        "Lexend",
    ]

    current_font = css_family_to_font_label(
        normalize_font_family(st.session_state.get("fontsel", "Avenir"))
    )
    try:
        default_index = font_options.index(current_font)
    except ValueError:
        default_index = font_options.index("Avenir")

    chosen_label = st.selectbox(
        "Choose a font for the UI and exported HTML:",
        options=font_options,
        index=default_index,
        key="fontsel_label",
    )


    # Map UI label to actual CSS font-family name used by @font-face.
    # OpenDyslexic must be the exact family name used in the generated CSS.
    fontsel = normalize_font_family(font_label_to_css_family(chosen_label))

    st.session_state.fontsel = fontsel

    # Content type selection: Book vs Script
    content_type = st.radio(
        "Content type:",
        options=["Book", "Script"],
        index=0 if st.session_state.get("content_type", "Book") == "Book" else 1,
        horizontal=True,
    )
    st.session_state.content_type = content_type

    # Existing user key input
    user_input = st.text_input(
        "Enter your user key (username, nickname, or passphrase):",
        key="userkey_input",
    )
    if st.button("Next"):
        if user_input.strip() == "":
            st.warning("You must enter a user key to continue.")
            st.stop()
        else:
            st.session_state.userkey = user_input.strip()
            st.session_state.step = 1
            st.rerun()
    st.stop()
    
def get_saved_colors_file():
    return f"{st.session_state.userkey}-speaker_colors.json"
    
def get_progress_file():
    return f"{st.session_state.userkey}-progress.json"
    
def get_unmatched_quotes_filename():
    return f"{st.session_state.userkey}-unmatched_quotes.txt"

def normalize_text(text):
    text = text.replace("\u00A0", " ")
    text = text.replace("…", "...")
    text = text.replace("“", "\"").replace("”", "\"")
    text = text.replace("’", "'").replace("‘", "'")
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def match_normalize(text):
    return text.replace("’", "'").replace("‘", "'")


def _is_boundary_char(ch: str) -> bool:
    return not ch or not ch.isalnum()


def find_with_boundaries_ci(haystack: str, needle: str, start: int = 0):
    """Case-insensitive, boundary-aware substring search."""
    if haystack is None or needle is None:
        return None
    h = haystack.lower()
    n = needle.lower()
    if not n:
        return None
    pos = h.find(n, max(0, start))
    while pos != -1:
        left_ok = (pos == 0) or _is_boundary_char(h[pos - 1])
        right_pos = pos + len(n)
        right_ok = (right_pos == len(h)) or _is_boundary_char(h[right_pos])
        if left_ok and right_ok:
            return pos
        pos = h.find(n, pos + 1)
    return None


def count_with_boundaries_ci(haystack: str, needle: str) -> int:
    if not haystack or not needle:
        return 0
    count = 0
    start = 0
    while True:
        pos = find_with_boundaries_ci(haystack, needle, start)
        if pos is None:
            return count
        count += 1
        start = pos + len(needle)


def nth_span_with_boundaries_ci(haystack: str, needle: str, target: int):
    """Return (start, end) for the target-th boundary-aware CI match."""
    if target <= 0:
        return None
    start = 0
    seen = 0
    while True:
        pos = find_with_boundaries_ci(haystack, needle, start)
        if pos is None:
            return None
        seen += 1
        if seen == target:
            return pos, pos + len(needle)
        start = pos + len(needle)


def normalize_text_with_index_map(text: str):
    """Normalize text while preserving a normalized-char -> raw-index map."""
    raw = str(text or "")
    out_chars = []
    out_map = []
    prev_space = True

    def emit(chars: str, raw_index: int):
        nonlocal prev_space
        for c in chars:
            if c.isspace():
                if prev_space:
                    continue
                out_chars.append(" ")
                out_map.append(raw_index)
                prev_space = True
            else:
                out_chars.append(c)
                out_map.append(raw_index)
                prev_space = False

    for i, ch in enumerate(raw):
        if ch == "\u00A0":
            emit(" ", i)
        elif ch == "…":
            emit("...", i)
        elif ch in ("“", "”"):
            emit('"', i)
        elif ch in ("’", "‘"):
            emit("'", i)
        else:
            emit(ch, i)

    # strip leading/trailing spaces in normalized text (to mirror normalize_text)
    start = 0
    end = len(out_chars)
    while start < end and out_chars[start] == " ":
        start += 1
    while end > start and out_chars[end - 1] == " ":
        end -= 1
    return "".join(out_chars[start:end]), out_map[start:end]


def extract_first_quoted_segment(text: str) -> str:
    """Return the first quoted segment if present, else the full text."""
    if text is None:
        return ""
    s = str(text)
    m = re.search(r'[“"]([^”"]+)[”"]', s)
    if not m:
        m = re.search(r"[‘']([^’']+)[’']", s)
    return m.group(1) if m else s


def compute_occurrence_target_for_review(quotes_records: list[dict], review_index: int) -> int:
    """Count how many earlier quote records carry the same normalized dialogue segment."""
    if not quotes_records or review_index is None or review_index < 0 or review_index >= len(quotes_records):
        return 1

    current = quotes_records[review_index] or {}
    current_text = (current.get("quote_with_marks") or current.get("quote_text") or "").strip()
    current_norm = normalize_text(extract_first_quoted_segment(current_text)).lower()
    if not current_norm:
        return 1

    prior_same = 0
    for i in range(review_index):
        rec = quotes_records[i] or {}
        prior_text = (rec.get("quote_with_marks") or rec.get("quote_text") or "").strip()
        prior_norm = normalize_text(extract_first_quoted_segment(prior_text)).lower()
        if prior_norm == current_norm:
            prior_same += 1

    return 1 + prior_same


def compute_start_paragraph_index_for_review(quotes_records: list[dict], review_index: int) -> int:
    """Use prior resolved context to keep Step 2 paragraph lookup moving forward."""
    if not quotes_records or review_index is None or review_index <= 0:
        return 0

    last_seen = 0
    for i in range(review_index):
        rec = quotes_records[i] or {}
        pidx = rec.get("paragraph_index")
        if isinstance(pidx, int) and pidx >= 0:
            last_seen = pidx
    return last_seen


def populate_record_context_fields(record: dict, context: dict, occurrence_target: int):
    """Write context lookup outputs into a quote record using the standard schema."""
    record["occurrence_target"] = occurrence_target
    record["paragraph_index"] = context.get("paragraph_index") if context else record.get("paragraph_index")
    record["context_previous_html"] = context.get("previous") if context else None
    record["context_current_html"] = context.get("current") if context else None
    record["context_next_html"] = context.get("next") if context else None
    record["context_previous_text"] = BeautifulSoup(context["previous"], "html.parser").get_text() if context and context.get("previous") else None
    record["context_current_text"] = BeautifulSoup(context["current"], "html.parser").get_text() if context and context.get("current") else None
    record["context_next_text"] = BeautifulSoup(context["next"], "html.parser").get_text() if context and context.get("next") else None


def autopopulate_context_for_all_records(quotes_records: list[dict]):
    """Populate paragraph/context fields for every quote record in sequence."""
    if not quotes_records:
        return {"updated": 0, "with_context": 0, "without_context": 0}

    with_context = 0
    without_context = 0
    for i, rec in enumerate(quotes_records):
        rec = rec or {}
        dialogue = (rec.get("quote_with_marks") or rec.get("quote_text") or "").strip()
        occurrence_target = compute_occurrence_target_for_review(quotes_records, i)
        start_paragraph_index = compute_start_paragraph_index_for_review(quotes_records, i)
        context = get_context_for_dialogue_json_only(
            dialogue,
            occurrence_target=occurrence_target,
            start_paragraph_index=start_paragraph_index,
        )
        populate_record_context_fields(rec, context, occurrence_target)
        if context:
            with_context += 1
        else:
            without_context += 1

    return {"updated": len(quotes_records), "with_context": with_context, "without_context": without_context}


def normalize_speaker_name(name):
    # Replace typographic apostrophes with straight ones, remove periods, lowercase, and trim.
    return name.replace("’", "'").replace("‘", "'").replace(".", "").lower().strip()

def smart_title(name):
    words = name.split()
    if not words:
        return name
    exceptions = {"ps", "pc", "ds", "di", "dci"}
    new_words = []
    for w in words:
        if w.lower() in exceptions:
            new_words.append(w.upper())
        else:
            new_words.append(w.capitalize())
    result = " ".join(new_words)
    result = re.sub(r"\(([mf])\)$", lambda m: "(" + m.group(1).upper() + ")", result, flags=re.IGNORECASE)
    return result


QUOTE_LINE_PATTERN = re.compile(r"^\s*(\d+(?:[a-zA-Z]+)?)\.\s+([^:]+):\s*(.*)$")


def parse_quote_line(raw_line: str):
    if raw_line is None:
        return None
    line = str(raw_line).rstrip("\n")
    m = QUOTE_LINE_PATTERN.match(line)
    if not m:
        return None
    idx_raw, speaker_raw, quote_raw = m.groups()
    try:
        index_num = int(re.match(r"^\d+", idx_raw).group(0))
    except Exception:
        index_num = None
    return {
        "index_raw": idx_raw,
        "index": index_num,
        "speaker_text": speaker_raw.strip(),
        "quote_text": quote_raw.strip(),
        "quote_with_marks": quote_raw.strip(),
    }


def make_quote_record(index: int, speaker_text: str, quote_text: str, quote_with_marks: str = None, content_type: str = None, index_raw: str = None):
    idx_raw = str(index_raw if index_raw is not None else index)
    return {
        "quote_id": f"q{index:05d}",
        "index": int(index),
        "index_raw": idx_raw,
        "speaker_text": (speaker_text or "Unknown").strip() or "Unknown",
        "quote_text": (quote_text or "").strip(),
        "quote_with_marks": (quote_with_marks if quote_with_marks is not None else quote_text or "").strip(),
        "review_status": "unreviewed",
        "predicted_speaker": None,
        "prediction_confidence": None,
        "candidate_speakers": [],
        "candidate_scores": {},
        "model_version": None,
        "content_type": content_type or st.session_state.get("content_type", "Book"),
        "paragraph_index": None,
        "occurrence_target": None,
        "context_previous_html": None,
        "context_current_html": None,
        "context_next_html": None,
        "context_previous_text": None,
        "context_current_text": None,
        "context_next_text": None,
        "notes": "",
        "manual_override": False,
    }


def normalize_record_schema(record: dict, index_fallback: int):
    raw_idx = record.get("index_raw")
    if raw_idx is None or str(raw_idx).strip() == "":
        raw_idx = str(record.get("index") or index_fallback)
    base = make_quote_record(
        index=int(record.get("index") or index_fallback),
        index_raw=raw_idx,
        speaker_text=record.get("speaker_text") or "Unknown",
        quote_text=record.get("quote_text") or record.get("quote_with_marks") or "",
        quote_with_marks=record.get("quote_with_marks") or record.get("quote_text") or "",
        content_type=record.get("content_type") or st.session_state.get("content_type", "Book"),
    )
    base.update(record or {})
    base["index"] = int(base.get("index") or index_fallback)
    base["index_raw"] = str(base.get("index_raw") or base["index"])
    base["quote_id"] = str(base.get("quote_id") or f"q{base['index']:05d}")
    if not isinstance(base.get("candidate_speakers"), list):
        base["candidate_speakers"] = []
    if not isinstance(base.get("candidate_scores"), dict):
        base["candidate_scores"] = {}
    return base


def record_to_legacy_line(record: dict) -> str:
    idx = record.get("index_raw") or record.get("index") or 0
    speaker = record.get("speaker_text") or "Unknown"
    quote_with_marks = record.get("quote_with_marks")
    if quote_with_marks is None or quote_with_marks == "":
        quote_with_marks = record.get("quote_text") or ""
    return f"{idx}. {speaker}: {quote_with_marks}\n"


def build_quotes_records_from_dialogue_list(dialogue_list):
    records = []
    content_type = st.session_state.get("content_type", "Book")
    for i, line in enumerate(dialogue_list or [], start=1):
        parsed = parse_quote_line(line)
        if parsed:
            rec = make_quote_record(
                index=parsed.get("index") or i,
                index_raw=parsed.get("index_raw") or str(parsed.get("index") or i),
                speaker_text=parsed.get("speaker_text") or "Unknown",
                quote_text=parsed.get("quote_text") or "",
                quote_with_marks=parsed.get("quote_with_marks") or "",
                content_type=content_type,
            )
        else:
            rec = make_quote_record(index=i, speaker_text="Unknown", quote_text=str(line).strip(), content_type=content_type)
        rec["quote_id"] = f"q{i:05d}"
        records.append(rec)
    return records


def build_quotes_records_from_quotes_lines(quotes_lines):
    dialogue = [str(line).rstrip("\n") for line in (quotes_lines or []) if str(line).strip()]
    return build_quotes_records_from_dialogue_list(dialogue)


def build_quotes_lines_from_records(quotes_records):
    return [record_to_legacy_line(r) for r in (quotes_records or [])]


def predict_speaker_for_record(record, quotes_records, session_state):
    """Future model seam. TODO: wire supervised attribution model output here."""
    return None


def get_next_record_for_review(quotes_records, start_index=0):
    if not quotes_records:
        return None, None
    for i in range(max(0, int(start_index)), len(quotes_records)):
        rec = quotes_records[i]
        speaker_unknown = (rec.get("speaker_text") or "").strip().lower() == "unknown"
        if speaker_unknown:
            return i, rec
    return None, None


def update_record_speaker(record: dict, new_speaker: str, action_type: str = "correct"):
    prev = record.get("speaker_text") or "Unknown"
    cleaned = (new_speaker or "").strip()
    if action_type == "skip":
        record["review_status"] = "skipped"
        return prev, prev
    if action_type == "ambiguous":
        record["review_status"] = "ambiguous"
        return prev, prev
    if not cleaned:
        return prev, prev
    cleaned = smart_title(cleaned)
    record["speaker_text"] = cleaned
    record["manual_override"] = True
    if prev.strip().lower() == "unknown":
        record["review_status"] = "corrected"
    elif normalize_speaker_name(prev) == normalize_speaker_name(cleaned):
        record["review_status"] = "accepted"
    else:
        record["review_status"] = "corrected"
    return prev, cleaned


def append_review_event(record: dict, action_type: str, previous_speaker: str, new_speaker: str):
    if "review_events" not in st.session_state or st.session_state.review_events is None:
        st.session_state.review_events = []
    st.session_state.review_events.append({
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "quote_id": record.get("quote_id"),
        "previous_speaker": previous_speaker,
        "new_speaker": new_speaker,
        "action_type": action_type,
        "review_status": record.get("review_status"),
        "predicted_speaker_at_action": record.get("predicted_speaker"),
        "prediction_confidence_at_action": record.get("prediction_confidence"),
    })


def sync_quotes_lines_from_records():
    st.session_state.quotes_lines = build_quotes_lines_from_records(st.session_state.get("quotes_records") or [])


def ensure_quotes_records_in_session():
    records = st.session_state.get("quotes_records")
    if records:
        st.session_state.quotes_records = [normalize_record_schema(r, i + 1) for i, r in enumerate(records)]
        sync_quotes_lines_from_records()
        return
    qlines = st.session_state.get("quotes_lines")
    if qlines:
        st.session_state.quotes_records = build_quotes_records_from_quotes_lines(qlines)
        sync_quotes_lines_from_records()


def migrate_legacy_state_to_quotes_records(data: dict):
    records = data.get("quotes_records")
    if records:
        return [normalize_record_schema(r, i + 1) for i, r in enumerate(records)]
    legacy_lines = data.get("quotes_lines")
    if legacy_lines:
        return build_quotes_records_from_quotes_lines(legacy_lines)
    return []

#def write_file_atomic(filepath, lines):
#    with open(filepath, "w", encoding="utf-8") as f:
#        f.writelines(lines)
#        f.flush()
#        os.fsync(f.fileno())

# ---------------------------
# Auto-Save & Auto-Load Functions
# ---------------------------
def auto_save():
    ensure_quotes_records_in_session()
    data = {
        "step": st.session_state.get("step", 1),
        "quotes_records": st.session_state.get("quotes_records"),
        "quotes_lines": st.session_state.get("quotes_lines"),
        "speaker_colors": st.session_state.get("speaker_colors"),
        "unknown_index": st.session_state.get("unknown_index", 0),
        "console_log": st.session_state.get("console_log", []),
        "review_events": st.session_state.get("review_events", []),
        "canonical_map": st.session_state.get("canonical_map") or {},
        "book_name": st.session_state.get("book_name"),
        "existing_speaker_colors": st.session_state.get("existing_speaker_colors"),
        "content_type": st.session_state.get("content_type", "Book")
    }
    if "docx_bytes" in st.session_state and st.session_state.docx_bytes is not None:
        data["docx_bytes"] = base64.b64encode(st.session_state.docx_bytes).decode("utf-8")
    with open(get_progress_file(), "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)
    if st.session_state.get("speaker_colors") is not None:
        with open(get_saved_colors_file(), "w", encoding="utf-8") as f:
            json.dump(st.session_state.speaker_colors, f, indent=4, ensure_ascii=False)
    if st.session_state.get("quotes_lines") and st.session_state.get("book_name"):
        quotes_filename = f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt"
        with open(quotes_filename, "w", encoding="utf-8") as f:
            quotes_text = "".join(st.session_state.quotes_lines)
            f.write(quotes_text)

def auto_load():
    if os.path.exists(get_progress_file()):
        with open(get_progress_file(), "r", encoding="utf-8") as f:
            data = json.load(f)
        data["quotes_records"] = migrate_legacy_state_to_quotes_records(data)
        if data.get("quotes_records"):
            data["quotes_lines"] = build_quotes_lines_from_records(data.get("quotes_records"))
        for key, value in data.items():
            st.session_state[key] = value
            # Normalise restored structures
            if isinstance(st.session_state.get("flagged_names"), list):
                st.session_state.flagged_names = set(st.session_state.flagged_names)
            if st.session_state.get("speaker_counts") is None:
                st.session_state.speaker_counts = {}
            if st.session_state.get("flagged_names") is None:
                st.session_state.flagged_names = set()
            if st.session_state.get("canonical_map") is None:
                st.session_state.canonical_map = {}
            if st.session_state.get("review_events") is None:
                st.session_state.review_events = []

            # Rebuild counts/flags from quotes_lines if missing or empty
            needs_rebuild = (
                not st.session_state.speaker_counts or
                (not st.session_state.flagged_names and st.session_state.speaker_counts)
            )
            if needs_rebuild and st.session_state.get("quotes_lines"):
                pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                counts_cap10 = {}
                flagged = set()
                for _line in st.session_state.quotes_lines:
                    m = pattern_speaker.match(_line.strip())
                    if not m:
                        continue
                    speaker_raw = m.group(1).strip()
                    effective = smart_title(speaker_raw)
                    norm = normalize_speaker_name(effective)
                    if norm in flagged:
                        continue
                    c = counts_cap10.get(norm, 0)
                    if c < 10:
                        c += 1
                        counts_cap10[norm] = c
                        if c >= 10:
                            flagged.add(norm)
                st.session_state.speaker_counts = counts_cap10
                st.session_state.flagged_names = flagged

        ensure_quotes_records_in_session()

        if "existing_speaker_colors" in st.session_state and st.session_state.existing_speaker_colors:
            st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in st.session_state.existing_speaker_colors.items()}
        if "docx_bytes" in st.session_state:
            docx_bytes = base64.b64decode(st.session_state["docx_bytes"].encode("utf-8"))
            st.session_state.docx_bytes = docx_bytes
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                tmp_docx.write(docx_bytes)
                st.session_state.docx_path = tmp_docx.name
            # Ensure d_json_path points to the unified JSON cache: [userkey]-[book_name].json in CWD
            try:
                userkey = st.session_state.get("userkey")
                book_name = st.session_state.get("book_name")
                if userkey and book_name:
                    cand = os.path.join(os.getcwd(), f"{userkey}-{book_name}.json")
                    if os.path.exists(cand):
                        st.session_state['d_json_path'] = cand
                    else:
                        # If the JSON cache is missing but we have a docx, rebuild it once
                        if st.session_state.get('docx_path') and os.path.exists(st.session_state['docx_path']):
                            write_paragraph_json_for_session()
            except Exception:
                pass


if os.path.exists(get_progress_file()):
    if st.button("Load Saved Progress"):
         auto_load()
         st.rerun()

# ---------------------------
# Alternative Extraction Functions
# ---------------------------
ATTACH_NO_SPACE = {"'", "’", "‘", '"', "“", "”", ",", ".", ";", ":", "?", "!"}
DASHES = {"-", "–", "—"}

def smart_join(run_texts):
    if not run_texts:
        return ""
    result = run_texts[0]
    for text in run_texts[1:]:
        if not text:
            continue

        # preserve any explicit/trailing spaces already in the text
        if result and result[-1].isspace():
            result += text.lstrip()
            continue
        if text[0].isspace():
            result += text
            continue

        # ellipses should attach to the previous token
        if text.startswith("...") or text.startswith("…"):
            result = result.rstrip()
            result += text
            continue

        prev = result[-1]
        first = text[0]

        # 1) Contractions/possessives: apostrophe binds to following letters (That’s, you’re, I’ll)
        if prev in {"'", "’", "‘"} and first.isalnum():
            result += text
            continue

        # 2) Characters that attach to the previous token (no leading space)
        if first in ATTACH_NO_SPACE:
            result += text
            continue

        # 3) Dashes/hyphens: attach tightly on both sides
        if first in DASHES:
            result += text            # no space before dash
            continue
        if prev in DASHES:
            result += text            # no space after dash
            continue

        # 4) Default spacing rule
        if prev.isalnum() and first.isalnum():
            result += text            # join words without extra space
        else:
            # Added guards (surgical): avoid space before ™/® and inside parentheses
            prev = result[-1] if result else ''
            first = text[0] if text else ''
            if first in {'™','®'} or prev == '(' or first == ')':
                result += text
                continue
            # Guard: keep opening double-quote tight with the next token
            prev = result[-1] if result else ''
            if prev in {'“', '"'}:
                result += text
                continue
            result += " " + text       # otherwise, insert a space
    return result

# def is_run_italic(run):
#    """Return True if the run is italic due to direct formatting or its character style."""
#    try:
#        if getattr(run.font, "italic", None) is True:
#            return True
#        style = getattr(run, "style", None)
#        if style is not None and getattr(style.font, "italic", None) is True:
#            return True
#    except Exception:
#        # Be conservative; if anything goes wrong, treat as non-italic
#        return False
#    return False
def effective_run_italic(run, paragraph):
    """Return True if, after cascading styles, this run is italic.
    Precedence (lowest to highest): paragraph style -> run character style -> direct run formatting.
    Explicit False overrides inherited True.
    """
    base = False
    try:
        # Paragraph style
        psty = getattr(paragraph, "style", None)
        if psty is not None and getattr(psty.font, "italic", None) is True:
            base = True
        # Character style on run
        rsty = getattr(run, "style", None)
        rsty_italic = getattr(getattr(rsty, "font", None), "italic", None)
        if rsty_italic is not None:
            base = bool(rsty_italic)
        # Direct run formatting
        rfmt_italic = getattr(getattr(run, "font", None), "italic", None)
        if rfmt_italic is not None:
            base = bool(rfmt_italic)
    except Exception:
        pass
    return base


def extract_italicized_text(paragraph):
    """
    def _trim_quote_edges(s: str) -> str:
        # Italics-path parity: remove spaces just inside opening/closing double quotes
        s = re.sub(r'(?<=[“\"])\\s+', '', s)
        s = re.sub(r'\\s+(?=[”\"])', '', s)
        return s
    Return a list of italic blocks for a paragraph.
    Detects italics after cascading: paragraph style -> character style -> direct run formatting.
    Preserves the existing >= 2-word threshold and smart_join behaviour.
    """
    italic_blocks = []
    current_block = []
    for run in paragraph.runs:
        if effective_run_italic(run, paragraph):
            current_block.append(run.text)
        else:
            joined = smart_join(current_block)
            joined = _trim_quote_edges(joined)
            raw_block = "".join(current_block)
            if len(joined.split()) >= 2 or is_single_titlecase_speaker_label(raw_block, next_text=""):
                italic_blocks.append(joined)
            current_block = []
    # flush tail
    joined = smart_join(current_block)
    joined = _trim_quote_edges(joined)
    joined = re.sub(r'^\.\s+(?=\w)', '', joined)
    raw_block = "".join(current_block)
    if len(joined.split()) >= 2 or is_single_titlecase_speaker_label(raw_block, next_text=""):
        italic_blocks.append(joined)
    return italic_blocks

def extract_italic_spans(paragraph):
    """
    Return a list of ((start, end), text) for contiguous italic blocks in this paragraph,
    using the same cascade logic and >=2-word rule as extract_italicized_text.
    Spans are computed against the raw concatenation of run.text (paragraph.text),
    with adjustment if a leading ". " is stripped.
    """
    spans = []
    pos = 0
    block_start = None
    block_raw = []  # raw run.text pieces

    for run in paragraph.runs:
        t = run.text or ""
        n = len(t)
        if effective_run_italic(run, paragraph):
            if block_start is None:
                block_start = pos
            block_raw.append(t)
        else:
            if block_start is not None:
                raw = "".join(block_raw)
                joined = smart_join(block_raw)
                shift = 0
                if re.match(r'^\.\s+(?=\w)', joined):
                    shift = 2
                    joined = joined[2:]
                next_text = paragraph.text[block_start + len(raw):block_start + len(raw) + 2] if (block_start + len(raw)) < len(paragraph.text) else ''
                if len(joined.split()) >= 2 or is_single_titlecase_speaker_label(raw[shift:], next_text=next_text):
                    start = block_start + shift
                    end = start + max(0, len(raw) - shift)
                                        # If this is a single-word speaker label across DOCX runs, ensure we output the colon.
                    if is_single_titlecase_speaker_label(raw[shift:], next_text=next_text):
                        label = raw[shift:].replace('\u00A0',' ').strip()
                        # If the colon isn't part of the italic run, it will be in next_text (e.g. ': ').
                        if not label.endswith(':') and str(next_text).startswith(':'):
                            label = label + ':'
                        joined = label
                    else:
                        joined = joined.strip()
                    spans.append(((start, end), joined))
                block_start = None
                block_raw = []
        pos += n

    if block_start is not None:
        raw = "".join(block_raw)
        joined = smart_join(block_raw)
        shift = 0
        if re.match(r'^\.\s+(?=\w)', joined):
            shift = 2
            joined = joined[2:]
        next_text = paragraph.text[block_start + len(raw):block_start + len(raw) + 2] if (block_start + len(raw)) < len(paragraph.text) else ''
        if len(joined.split()) >= 2 or is_single_titlecase_speaker_label(raw[shift:], next_text=next_text):
            start = block_start + shift
            end = start + max(0, len(raw) - shift)
            spans.append(((start, end), joined))

    return spans

def is_all_caps_name(s: str) -> bool:
    """
    True if all alphabetic characters in s are uppercase and there is at least one letter.
    Punctuation, digits and spaces are ignored for the check.
    """
    letters = [c for c in s if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def parse_docx_script(docx_path: str):
    """
    Parse a DOCX script and return a list of {"speaker": ..., "text": ...} using three patterns:

      1) Name: Dialogue
         - First colon splits name vs dialogue.
         - Name may be any case.
         - Left side must *look* like a name (<= 4 words, each word with letters starts uppercase).
         - Right side must contain at least one 'dialogue-like' word:
             * has any lowercase letters, or
             * is a single-letter uppercase word (I, A, etc).

      2) NAME Dialogue
         - No colon.
         - First token is ALL CAPS (NAME).
         - If the char immediately after NAME is a TAB:
             * any non-empty remainder counts as dialogue (all caps allowed).
         - If it is a SPACE (or other non-tab whitespace):
             * look only at the FIRST word after NAME:
               - dialogue only if that word has lowercase letters, or
               - is a single-letter uppercase word (I, A, etc).
             * If not, this line is NOT a dialogue line (e.g. MUSIC TRANSITION).

      3) NAME
         Dialogue
         - Whole line is ALL CAPS → speaker cue.
         - Following non-blank lines that are not new cues become dialogue until a blank or new cue.

    One DOCX paragraph is treated as one line.
    """
    import docx

    doc = docx.Document(docx_path)
    lines = [p.text.rstrip("\n") for p in doc.paragraphs]

    results = []
    current_speaker = None
    current_lines = []

    def flush():
        nonlocal current_speaker, current_lines
        if current_speaker and current_lines:
            text = " ".join(t.strip() for t in current_lines if t.strip())
            if text:
                results.append({"speaker": current_speaker.strip(), "text": text})
        current_speaker = None
        current_lines = []

    def has_letters(w: str) -> bool:
        return any(c.isalpha() for c in w)

    def dialogue_word_anywhere(word: str) -> bool:
        """
        'Dialogue-like' word:
          - has any lowercase letters; OR
          - is a single-letter uppercase word (I, A, etc).
        """
        alpha = [c for c in word if c.isalpha()]
        if not alpha:
            return False
        if any(c.islower() for c in alpha):
            return True
        return len(alpha) == 1 and alpha[0].isupper()

    for raw in lines:
        line = raw.rstrip("\r\n")
        s = line.strip()

        # Blank line ends current block
        if not s:
            flush()
            continue

        # ---------- Pattern 1: Name: Dialogue ----------
        if ":" in s:
            before, after = s.split(":", 1)
            name_part = before.strip()
            rest = after.lstrip()

            if name_part and rest:
                name_tokens = name_part.split()
                # "Looks like a name" = 1–4 tokens, each token with letters starts uppercase
                name_tokens_with_letters = [t for t in name_tokens if has_letters(t)]
                looks_like_name = (
                    len(name_tokens_with_letters) > 0
                    and len(name_tokens) <= 4
                    and all(t[0].isupper() for t in name_tokens_with_letters)
                )

                words_rest = rest.split()
                dialogue_exists = any(dialogue_word_anywhere(w) for w in words_rest)

                if looks_like_name and dialogue_exists:
                    flush()
                    current_speaker = name_part
                    current_lines = [rest]
                    continue

            # Not a cue; maybe continuation text
            if current_speaker:
                current_lines.append(s)
            continue

        # ---------- No colon: patterns 2 and 3 ----------

        line_stripped = s
        tokens = line_stripped.split()

        if tokens:
            first_tok = tokens[0]

            # Pattern 2: NAME Dialogue (no colon, first token ALL CAPS)
            if is_all_caps_name(first_tok):
                # Character immediately after NAME in the stripped line
                delim_char = line_stripped[len(first_tok)] if len(line_stripped) > len(first_tok) else " "
                rest_str = line_stripped[len(first_tok):].lstrip()

                if rest_str:
                    if delim_char == "\t":
                        # NAME<TAB>Dialogue: allow all caps dialogue
                        flush()
                        current_speaker = first_tok
                        current_lines = [rest_str]
                        continue
                    else:
                        # NAME<space>Dialogue: check FIRST word only
                        rest_words = rest_str.split()
                        first_rest_word = rest_words[0] if rest_words else ""
                        if first_rest_word and dialogue_word_anywhere(first_rest_word):
                            flush()
                            current_speaker = first_tok
                            current_lines = [rest_str]
                            continue
                # If there's no remainder or it doesn't look like dialogue,
                # fall through to possible Pattern 3 (NAME alone) / continuation.

        # Pattern 3: NAME on its own line
        if is_all_caps_name(s):
            flush()
            current_speaker = s
            continue

        # Continuation of current speaker
        if current_speaker:
            current_lines.append(s)
        # Else: stage directions / SFX / headings are ignored

    flush()
    return results


def extract_dialogue_from_docx_script(docx_path: str):
    """
    Use parse_docx_script and return a list of numbered lines in the same logical
    format as quotes.txt:

        "1. Speaker: Dialogue"

    Speaker names are normalised with smart_title, dialogue text is left as-is.
    """
    pairs = parse_docx_script(docx_path)
    lines = []
    line_number = 1
    for pair in pairs:
        raw_speaker = (pair.get("speaker") or "").strip()
        text = (pair.get("text") or "").strip()
        if not raw_speaker or not text:
            continue

        speaker = smart_title(raw_speaker)  # JOHN HOLMES -> John Holmes, etc.
        lines.append(f"{line_number}. {speaker}: {text}")
        line_number += 1

    return lines


def extract_dialogue_from_docx(book_name, docx_path):
    # Helpers: italics-path check for quote enclosure (compare-only, no text mutation)
    import re as _re_local
    _SPACE_LIKE = _re_local.compile(r'[\u0020\u00A0\u2009\u200A\u200B\u202F\u205F\u3000]+')
    _OPEN_QS  = {'“', '"'}
    _CLOSE_QS = {'”', '"'}
    def _prev_non_space(_s: str, _idx: int) -> str:
        i = _idx - 1
        while i >= 0 and _SPACE_LIKE.match(_s[i]):
            i -= 1
        return _s[i] if i >= 0 else ''
    def _next_non_space(_s: str, _idx: int) -> str:
        n = len(_s); i = _idx
        while i < n and _SPACE_LIKE.match(_s[i]):
            i += 1
        return _s[i] if i < n else ''
    def _is_enclosed_by_quotes(_para_text: str, _start: int, _end: int, _seg_text: str) -> bool:
        # Pattern A: quotes outside the italic span: … “ [italic] ” …
        _left  = _prev_non_space(_para_text, _start)
        _right = _next_non_space(_para_text, _end)
        if _left in _OPEN_QS and _right in _CLOSE_QS:
            return True
        # Pattern B: quotes inside the italic span text (rare)
        st = _seg_text
        if st and st[0] in _OPEN_QS:
            j = 1
            while j < len(st) and _SPACE_LIKE.match(st[j]):
                j += 1
            st = st[:1] + st[j:]
        if st and st[-1] in _CLOSE_QS:
            k = len(st) - 2
            while k >= 0 and _SPACE_LIKE.match(st[k]):
                k -= 1
            st = st[:k+1] + st[-1:]
        return (len(st) >= 2 and st[0] in _OPEN_QS and st[-1] in _CLOSE_QS)
    def _split_interrupted_dialogue(seg_text: str):
        """
        Split a quoted dialogue segment when it contains a narration interruption
        wrapped in dashes, e.g.:
          “There can’t be. And yet”—she raised her head—“and yet sometimes ...”
        Returns one or more dialogue-only pieces (narration interruption removed).
        """
        s = (seg_text or "").strip()
        if not s:
            return []

        # Match: [optional open quote] left [close quote]—aside—[open quote] right [optional close quote]
        m = re.match(
            r'^\s*([“"]?)(.+?)([”"])\s*[—–]\s*([^—–]+?)\s*[—–]\s*([“"])(.+?)([”"]?)\s*$',
            s
        )
        if not m:
            return [s]

        open1, left, close1, aside, open2, right, close2 = m.groups()
        if not aside or not re.search(r'[A-Za-z]', aside):
            return [s]

        left_piece = f"{open1}{left.strip()}{close1}".strip()
        right_piece = f"{open2}{right.strip()}{close2}".strip()
        return [p for p in (left_piece, right_piece) if p]
    doc = docx.Document(docx_path)
    quote_pattern = re.compile(r'(?:^|\s)(["“].+?["”])(?=$|[\s\.\,\;\:\!\?\)\]\}])')
    dialogue_list = []
    line_number = 1
    for para in doc.paragraphs:
        text = para.text.strip()
    
        # Build ordered segments: closing-only -> paired -> opening-only
        matches = list(quote_pattern.finditer(text))  # paired matches
        covered = []

        def _overlaps(a, b):
            # a, b are (start, end) half-open intervals
            return not (a[1] <= b[0] or b[1] <= a[0])

        def _is_covered(span):
            return any(_overlaps(span, c) for c in covered)

        OPEN = {'“', '"'}
        CLOSE = {'”', '"'}

        ordered = []

        # 1) Closing-only: first closing before any opening in *uncovered* text
        first_close = -1
        first_open = -1
        for i, ch in enumerate(text):
            if _is_covered((i, i + 1)):
                continue
            if ch in CLOSE and first_close == -1:
                first_close = i
            if ch in OPEN and first_open == -1:
                first_open = i
            if first_close != -1 and (first_open == -1 or first_close < first_open):
                seg_span = (0, first_close + 1)
                seg = text[seg_span[0]:seg_span[1]].strip()
                if seg:
                    ordered.append((seg_span, seg))
                    covered.append(seg_span)
                break  # only the first closing-only segment per paragraph

        # 2) Paired quotes: use existing regex; skip spans already covered
        for m in matches:
            seg_span = (m.start(1), m.end(1))
            if not _is_covered(seg_span):
                seg = m.group(1).strip()
                if seg:
                    ordered.append((seg_span, seg))
                    covered.append(seg_span)

        # 3) Opening-only: last opening with no closing after -> opening..end
        last_open = -1
        for i, ch in enumerate(text):
            if ch in OPEN and not _is_covered((i, i + 1)):
                last_open = i

        if last_open != -1:
            # any *uncovered* closing after last_open?
            has_close_after = False
            j = last_open + 1
            while j < len(text):
                if not _is_covered((j, j + 1)) and text[j] in CLOSE:
                    has_close_after = True
                    break
                j += 1
            if not has_close_after:
                seg_span = (last_open, len(text))
                if not _is_covered(seg_span):
                    seg = text[seg_span[0]:seg_span[1]].strip()
                    if seg:
                        ordered.append((seg_span, seg))
                        covered.append(seg_span)

        # Merge quotes (with spans) and italics (with spans), then sort by reading order
        items = []  # list of ((start, end), text)

        # quotes collected earlier as (span, text)
        for span, seg in ordered:
            items.append((span, seg))

        quote_spans = [span for span, _ in ordered]

        def _inside_any(inner_span, outer_spans):
            s, e = inner_span
            return any(os <= s and e <= oe for (os, oe) in outer_spans)

        for span, seg in extract_italic_spans(para):
            # Skip italics that lie anywhere inside any quoted span in this paragraph
            if _inside_any(span, quote_spans):
                continue
            items.append((span, seg))

        items.sort(key=lambda it: (it[0][0], -(it[0][1] - it[0][0])))

        for _, seg in items:
            seg_clean = (seg or "").strip()
            if seg_clean:
                for seg_part in _split_interrupted_dialogue(seg_clean):
                    dialogue_list.append(f"{line_number}. Unknown: {seg_part}")
                    line_number += 1
    
                continue
            # If the segment is empty after stripping, skip it
            line_number += 1
            continue
            line_number += 1
    output_path = f"{st.session_state.userkey}-{book_name}-quotes.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(dialogue_list))
    return dialogue_list

# ---------------------------
# DOCX-to-HTML & Marking Functions
# ---------------------------

def prepend_marker_to_paragraph(paragraph, marker_text):
    p = paragraph._p
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = marker_text + " "
    r.append(t)
    p.insert(0, r)

def create_marker_docx(original_docx, marker_docx):
    doc = docx.Document(original_docx)
    for idx, para in enumerate(doc.paragraphs):
        marker = f"[[[P{idx}]]]"
        prepend_marker_to_paragraph(para, marker)
    doc.save(marker_docx)

def convert_docx_to_html_mammoth(docx_file):
    with open(docx_file, "rb") as f:
        result = mammoth.convert_to_html(f)
        return result.value

def get_manual_indentation(docx_file):
    doc = docx.Document(docx_file)
    indented_paras = {}
    for idx, para in enumerate(doc.paragraphs):
        left = para.paragraph_format.left_indent
        right = para.paragraph_format.right_indent
        if (left is not None and left.pt > 0) or (right is not None and right.pt > 0):
            indented_paras[idx] = (left, right)
    return indented_paras

def convert_length_to_px(length):
    return length.pt * 1.33 if length is not None else 0

# ---------------------------
# Dialogue Highlighting Functions
# ---------------------------
def highlight_across_nodes(parent, quote, highlight_style, soup):
    full_text = parent.get_text()
    full_text_lower = match_normalize(full_text).lower()
    quote_lower = match_normalize(quote).lower()
    idx = full_text_lower.find(quote_lower)
    if idx == -1:
        return False
    end_idx = idx + len(quote)
    running_index = 0
    for descendant in list(parent.descendants):
        if isinstance(descendant, NavigableString):
            text = str(descendant)
            text_length = len(text)
            node_start = running_index
            node_end = running_index + text_length
            if node_end > idx and node_start < end_idx:
                overlap_start = max(idx, node_start)
                overlap_end = min(end_idx, node_end)
                rel_start = overlap_start - node_start
                rel_end = overlap_end - node_start
                before = text[:rel_start]
                match_text = text[rel_start:rel_end]
                after = text[rel_end:]
                new_nodes = []
                if before:
                    new_nodes.append(NavigableString(before))
                span_tag = soup.new_tag("span", attrs={"class": "highlight", "style": highlight_style})
                span_tag.string = match_text
                new_nodes.append(span_tag)
                if after:
                    new_nodes.append(NavigableString(after))
                descendant.replace_with(*new_nodes)
            running_index += text_length
    return True

def highlight_quote_in_parent(parent, quote, highlight_style, soup):
    stripped_quote = quote.strip('“”"')
    stripped_quote_lower = match_normalize(stripped_quote).lower()
    for child in parent.contents:
        if isinstance(child, NavigableString):
            child_text = str(child)
            child_text_lower = match_normalize(child_text).lower()
            idx = child_text_lower.find(stripped_quote_lower)
            if idx != -1:
                before = child_text[:idx]
                match_text = child_text[idx: idx + len(stripped_quote)]
                after = child_text[idx + len(stripped_quote):]
                new_nodes = []
                if before:
                    new_nodes.append(NavigableString(before))
                span_tag = soup.new_tag("span", attrs={"class": "highlight", "style": highlight_style})
                span_tag.string = match_text
                new_nodes.append(span_tag)
                if after:
                    new_nodes.append(NavigableString(after))
                child.replace_with(*new_nodes)
                return True
        elif hasattr(child, 'contents'):
            if highlight_quote_in_parent(child, quote, highlight_style, soup):
                return True
    return highlight_across_nodes(parent, stripped_quote, highlight_style, soup)

def find_with_boundaries(haystack, needle, start=0):
    """Find needle in haystack from start, requiring word-boundaries when needle begins/ends with alphanumerics.

    This prevents matching inside larger words (e.g. matching 'or,' inside 'for,').
    """
    if not needle:
        return -1
    need_start_boundary = needle[0].isalnum()
    need_end_boundary = needle[-1].isalnum()

    pos = haystack.find(needle, start)
    while pos != -1:
        before_ok = True
        after_ok = True

        if need_start_boundary:
            before_ok = (pos == 0) or (not haystack[pos - 1].isalnum())
        if need_end_boundary:
            after_idx = pos + len(needle)
            after_ok = (after_idx >= len(haystack)) or (not haystack[after_idx].isalnum())

        if before_ok and after_ok:
            return pos

        pos = haystack.find(needle, pos + 1)

    return -1

def build_candidate_info(soup):
    candidates = soup.find_all(['p', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
    candidate_info = []
    global_offset = 0
    for candidate in candidates:
        text = candidate.get_text()
        length = len(text)
        candidate_info.append((candidate, global_offset, global_offset + length, text))
        global_offset += length
    return candidate_info

def highlight_in_candidate(candidate, quote, highlight_style, soup, start_offset=0, strict=False):
    full_text = candidate.get_text()

    # Case-sensitive, boundary-aware search (strictness is controlled by what 'quote' is:
    # e.g. quote_with_marks vs quote without marks).
    pos = find_with_boundaries(full_text, quote, start_offset)
    if pos == -1:
        return None
    match_end = pos + len(quote)

    running_index = 0
    for descendant in list(candidate.descendants):
        if isinstance(descendant, NavigableString):
            text = str(descendant)
            text_length = len(text)
            node_start = running_index
            node_end = running_index + text_length

            overlap_start = max(pos, node_start)
            overlap_end = min(match_end, node_end)

            if overlap_start < overlap_end:
                rel_start = overlap_start - node_start
                rel_end = overlap_end - node_start
                before = text[:rel_start]
                match_text = text[rel_start:rel_end]
                after = text[rel_end:]
                new_nodes = []
                if before:
                    new_nodes.append(NavigableString(before))
                span_tag = soup.new_tag("span", attrs={"class": "highlight", "style": highlight_style})
                span_tag.string = match_text
                new_nodes.append(span_tag)
                if after:
                    new_nodes.append(NavigableString(after))
                descendant.replace_with(*new_nodes)

            running_index += text_length

    return match_end

def highlight_dialogue_in_html(html, quotes_list, speaker_colors):
    soup = BeautifulSoup(html, "html.parser")
    candidate_info = build_candidate_info(soup)
    unmatched_quotes = []
    last_global_offset = 0

    def style_for_speaker(speaker):
        norm_speaker = normalize_speaker_name(speaker)
        color_choice = st.session_state.speaker_colors.get(norm_speaker, "none")
        if norm_speaker == "unknown":
            color_choice = "none"
        rgba = COLOR_PALETTE.get(color_choice, COLOR_PALETTE["none"])
        if color_choice == "none":
            return f"color: rgb({rgba[0]}, {rgba[1]}, {rgba[2]}); background-color: transparent;"
        return f"color: {rgba[4]}; background-color: rgba({rgba[0]}, {rgba[1]}, {rgba[2]}, {rgba[3]});"

    def search_and_highlight_from_global(needle, start_global):
        """Search forward from a global offset for needle (case-sensitive, boundary-aware)."""
        nonlocal last_global_offset
        if not needle:
            return False

        for candidate, start, end, text in candidate_info:
            if end <= start_global:
                continue

            local_start = max(0, start_global - start) if start_global > start else 0
            pos = find_with_boundaries(text, needle, local_start)
            if pos == -1:
                continue

            match_end_local = highlight_in_candidate(candidate, needle, current_style, soup, local_start, strict=True)
            if match_end_local is None:
                continue

            last_global_offset = start + match_end_local
            return True

        return False

    for quote_data in quotes_list:
        speaker = quote_data.get("speaker", "")
        current_style = style_for_speaker(speaker)

        quote_with_marks = (quote_data.get("quote_with_marks") or "").strip()
        quote_plain = (quote_data.get("quote") or "").strip()

        matched = False

        # Stage 1: from current position, search forwards for the quote INCLUDING quote marks (case-sensitive)
        if quote_with_marks:
            matched = search_and_highlight_from_global(quote_with_marks, last_global_offset)

        # Stage 2: if not found, search again from the start for the quote INCLUDING quote marks (case-sensitive)
        if (not matched) and quote_with_marks:
            matched = search_and_highlight_from_global(quote_with_marks, 0)

        # Stage 3: if still not found, search from the start WITHOUT quote marks (case-sensitive)
        if (not matched) and quote_plain:
            matched = search_and_highlight_from_global(quote_plain, 0)

        # Stage 4: if still not found, mark as unmatched (and keep prior behaviour: save to file)
        if not matched:
            unmatched_quotes.append(f"{quote_data.get('speaker','')}: \"{quote_data.get('quote','')}\" [Index: {quote_data.get('index','')}]")

    if unmatched_quotes:
        unmatched_quotes_filename = get_unmatched_quotes_filename()
        with open(unmatched_quotes_filename, "w", encoding="utf-8") as f:
            f.write("\n".join(unmatched_quotes))
        st.write(f"⚠️ Unmatched quotes saved to '[userkey]-unmatched_quotes.txt' ({len(unmatched_quotes)} entries)")

    return str(soup)

def apply_manual_indentation_with_markers(original_docx, html):
    indented_paras = get_manual_indentation(original_docx)
    soup = BeautifulSoup(html, "html.parser")
    marker_regex = re.compile(r"\[\[\[P(\d+)\]\]\]")
    candidate_tags = ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']
    for tag in soup.find_all(candidate_tags):
        text = tag.get_text()
        match = marker_regex.search(text)
        if match:
            para_index = int(match.group(1))
            for text_node in tag.find_all(string=True):
                if marker_regex.search(text_node):
                    new_text = marker_regex.sub("", text_node)
                    text_node.replace_with(new_text)
            if para_index in indented_paras:
                left, right = indented_paras[para_index]
                left_px = convert_length_to_px(left)
                right_px = convert_length_to_px(right)
                style_str = f"margin-left: {left_px}px; margin-right: {right_px}px;"
                if tag.has_attr("style"):
                    tag["style"] += " " + style_str
                else:
                    tag["style"] = style_str
    return str(soup)

def transform_script_layout(html: str) -> str:
    """
    Post-process the HTML produced by Mammoth + highlighting so that
    script dialogue lines are rendered as:

        <p class="script-line">
          <span class="script-speaker">NAME:</span>
          <span class="script-dialogue">...dialogue (with highlights)...</span>
        </p>

    We only transform <p> elements that:
      - contain at least one <span class="highlight"> (i.e. actual dialogue), and
      - start with something like 'NAME:' in ALL CAPS.
    """
    soup = BeautifulSoup(html, "html.parser")

    for p in soup.find_all("p"):
        # Only touch paragraphs that contain highlighted dialogue
        if not p.find("span", class_="highlight"):
            continue

        full_text = p.get_text()
        # Match leading ALLCAPS speaker name followed by a colon
        m = re.match(r"^\s*([A-Z][A-Z0-9 ]{0,50})\s*:\s*", full_text)
        if not m:
            continue

        speaker = m.group(1).strip()

        # Remove the speaker prefix from the *HTML* content, not just the text
        inner_html = p.decode_contents()
        # Strip "  NAME   :   " + optional tabs/spaces
        prefix_pattern = r"^\s*" + re.escape(speaker) + r"\s*:\s*[\t ]*"
        dialogue_html, n_subs = re.subn(prefix_pattern, "", inner_html, count=1)
        if n_subs == 0:
            # Couldn't safely strip; skip this paragraph
            continue

        # Clear the paragraph and rebuild
        p.clear()

        # Remove margin-left from inline style, keep other style properties
        if p.has_attr("style"):
            parts = [part.strip() for part in p["style"].split(";") if part.strip()]
            parts = [part for part in parts if not part.lower().startswith("margin-left")]
            if parts:
                p["style"] = "; ".join(parts)
            else:
                del p["style"]

        # Add a class for styling
        existing_classes = p.get("class", [])
        if "script-line" not in existing_classes:
            existing_classes.append("script-line")
        if existing_classes:
            p["class"] = existing_classes

        # Speaker span
        speaker_span = soup.new_tag("span", attrs={"class": "script-speaker"})
        speaker_span.string = speaker + ":"
        p.append(speaker_span)
        p.append(" ")

        # Dialogue span – preserve existing highlight spans etc.
        dialogue_span = soup.new_tag("span", attrs={"class": "script-dialogue"})
        frag = BeautifulSoup(dialogue_html, "html.parser")
        for child in frag.contents:
            dialogue_span.append(child)
        p.append(dialogue_span)

    return str(soup)


# -------------------------
# ---------------------------
# Summary & Ranking Functions
# ---------------------------
def generate_summary_html(quotes_list, speakers, speaker_colors):
    counts = Counter(quote["speaker"] for quote in quotes_list)
    total_lines = sum(counts.values())
    summary_order = []
    if "Unknown" in counts:
        summary_order.append("Unknown")
    for sp in speakers:
        if sp != "Unknown" and sp not in summary_order:
            summary_order.append(sp)
    for sp in counts:
        if sp not in summary_order:
            summary_order.append(sp)
    lines = []
    lines.append('<div id="character-summary" style="border: 1px solid #ccc; padding: 10px; margin-bottom: 20px;">')
    lines.append('<h2 style="margin: 0 0 5px 0;">Character Summary</h2>')
    for sp in summary_order:
        if normalize_speaker_name(sp) in ("error", "do not read"):
            continue
        count = counts.get(sp, 0)
        percentage = round((count / total_lines) * 100) if total_lines > 0 else 0
        color_key = speaker_colors.get(normalize_speaker_name(sp), "none")
        if sp.lower() == "unknown":
            color_key = "none"
        rgba = COLOR_PALETTE.get(color_key, COLOR_PALETTE["none"])
        if color_key == "none":
            style = f"color: rgb({rgba[0]}, {rgba[1]}, {rgba[2]}); background-color: transparent;"
        else:
            style = f"color: {rgba[4]}; background-color: rgba({rgba[0]}, {rgba[1]}, {rgba[2]}, {rgba[3]});"
        lines.append(f'<p style="margin: 0; line-height: 1.2; padding: 8px 0;"><span class="highlight" style="{style}">{sp}</span> - {count} lines - {percentage}%</p>')
    lines.append('</div>')
    return "\n".join(lines)

def generate_ranking_html(quotes_list, speaker_colors):
    counts = Counter(quote["speaker"] for quote in quotes_list)
    total_lines = sum(counts.values())
    filtered = [(sp, count) for sp, count in counts.items() if sp.lower() not in ("unknown", "do not read", "error") and count > 1]
    filtered.sort(key=lambda x: x[1], reverse=True)
    lines = []
    lines.append('<div id="speaker-ranking" style="margin-top: 20px;">')
    lines.append('<h2 style="margin: 0 0 5px 0;">Speaker Ranking</h2>')
    for sp, count in filtered:
        percentage = round((count / total_lines) * 100) if total_lines > 0 else 0
        color_key = speaker_colors.get(normalize_speaker_name(sp), "none")
        rgba = COLOR_PALETTE.get(color_key, COLOR_PALETTE["none"])
        if color_key == "none":
            style = f"color: rgb({rgba[0]}, {rgba[1]}, {rgba[2]}); background-color: transparent;"
        else:
            style = f"color: {rgba[4]}; background-color: rgba({rgba[0]}, {rgba[1]}, {rgba[2]}, {rgba[3]});"
        lines.append(f'<p style="margin: 0; line-height: 1.2; padding: 8px 0;"><span class="highlight" style="{style}">{sp}</span> - {count} lines - {percentage}%</p>')
    lines.append('</div>')
    return "\n".join(lines)

def generate_first_lines_html(quotes_list, speakers):
    # Map: speaker (canonical) -> first qualifying quote
    first_lines = {}
    for quote in quotes_list:
        speaker = quote["speaker"]
        norm = normalize_speaker_name(speaker)
        # Only add the first qualifying line (3+ words, else first)
        if norm not in first_lines:
            words = quote["quote"].strip().split()
            if len(words) >= 3 or all(q["speaker"] != speaker for q in quotes_list if q != quote):
                first_lines[norm] = quote["quote"].strip()
            else:
                # Tentatively store, may be replaced by a later 3+ word line
                first_lines[norm] = quote["quote"].strip()
        else:
            if len(first_lines[norm].split()) < 3 and len(quote["quote"].strip().split()) >= 3:
                first_lines[norm] = quote["quote"].strip()

    lines = []
    lines.append('<div id="first-lines-summary" style="border: 1px solid #ccc; padding: 10px; margin-bottom: 20px;">')
    lines.append('<h2 style="margin: 0 0 5px 0;">First Substantial Lines</h2>')
    for sp in speakers:
        norm = normalize_speaker_name(sp)
        if norm in ("do not read", "error", "unknown"):
            continue
        if norm in first_lines:
            line = first_lines[norm]
            lines.append(f'<p style="margin: 0; line-height: 1.2; padding: 8px 0;"><span class="highlight">{sp}</span>: <span style="font-style: italic;">{line}</span></p>')
    lines.append('</div>')
    return "\n".join(lines)
  
# ---------------------------
# Canonical Speaker & Quote Functions
# ---------------------------
def get_canonical_speakers(quotes_file):
    speakers = []
    pattern = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
    with open(quotes_file, "r", encoding="utf-8") as f:
        for line in f:
            match = pattern.match(line.strip())
            if match:
                speaker_raw = match.group(1).strip()
                speakers.append(smart_title(str(speaker_raw)))  # Ensure it's a string
    seen = set()
    canonical_speakers = []
    for s in speakers:
        norm = normalize_speaker_name(str(s))
        if norm not in seen:
            seen.add(norm)
            canonical_speakers.append(s)
    canonical_map = {normalize_speaker_name(str(s)): s for s in canonical_speakers}
    return canonical_speakers, canonical_map

def load_quotes(quotes_file, canonical_map):
    quotes_list = []
    # Capture optional opening/closing quotes so we can do a strict first-pass match INCLUDING quote marks.
    pattern = re.compile(r"^\s*([0-9]+(?:[a-zA-Z]+)?)\.\s+([^:]+):\s*([“\"])?(.+?)([”\"])?\s*$")
    with open(quotes_file, "r", encoding="utf-8") as f:
        for line in f:
            match = pattern.match(line.strip())
            if match:
                index, speaker_raw, open_q, quote_inner_raw, close_q = match.groups()
                effective = smart_title(speaker_raw)
                norm = normalize_speaker_name(effective)
                canonical = canonical_map.get(norm, effective)

                quote_inner = quote_inner_raw.strip()
                quote_with_marks = f"{open_q or ''}{quote_inner}{close_q or ''}"

                quotes_list.append({
                    "index": index,
                    "speaker": canonical,
                    "quote": quote_inner,
                    "quote_with_marks": quote_with_marks
                })
    return quotes_list

def load_existing_colors():
  if os.path.exists(get_saved_colors_file()):
    with open(get_saved_colors_file(), "r", encoding="utf-8") as f:
        loaded_colors = json.load(f)
    normalized_loaded = {normalize_speaker_name(k): v for k, v in loaded_colors.items()}
    st.session_state.speaker_colors = normalized_loaded
    st.session_state.existing_speaker_colors = normalized_loaded

def save_speaker_colors(speaker_colors):
    with open(get_saved_colors_file(), "w", encoding="utf-8") as f:
        json.dump(speaker_colors, f, indent=4, ensure_ascii=False)

# ---------------------------
# Restart Helper Function
# ---------------------------
def restart_app():
    st.session_state.clear()
    st.rerun()

# ---------------------------
# Streamlit Multi-Step UI
# ---------------------------

# ========= STEP 1: Upload & Initialize =========
if st.session_state.step == 1:
    ensure_quotes_records_in_session()
    st.markdown("<h4>DOCX to HTML Converter with Dialogue Highlighting</h4>", unsafe_allow_html=True)
    st.write("Upload your DOCX and quotes text files. Optionally, upload an existing speaker_colors.json file.")
    st.write("Alternatively, upload **just a DOCX** to create a quotes text file.")
    
    if "docx_bytes" in st.session_state:
        st.success("DOCX already uploaded and processed.")
        if st.session_state.get("docx_only", False):
            if st.session_state.get("quotes_lines") is not None:
                quotes_txt = "\n".join(st.session_state.quotes_lines)
                st.download_button("Download Extracted Quotes TXT", quotes_txt.encode("utf-8"),
                                   file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt", mime="text/plain")
                if st.button("Restart", key="restart_docx"):
                    restart_app()
                if st.button("Continue", key="continue_docx"):
                    st.session_state.docx_only = False
                    st.session_state.unknown_index = 0
                    st.session_state.console_log = []
                    if st.session_state.get("content_type", "Book") == "Script":
                        st.session_state.step = 3
                    else:
                        st.session_state.step = 2
                    # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
                    try:
                        if st.session_state.get("quotes_lines"):
                            pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                            counts_cap10 = {}
                            flagged = set()
                            for _line in st.session_state.quotes_lines:
                                m = pattern_speaker.match(_line.strip())
                                if not m:
                                    continue
                                speaker_raw = m.group(1).strip()
                                effective = smart_title(speaker_raw)
                                norm = normalize_speaker_name(effective)
                                if norm in flagged:
                                    continue
                                c = counts_cap10.get(norm, 0)
                                if c < 10:
                                    c += 1
                                    counts_cap10[norm] = c
                                    if c >= 10:
                                        flagged.add(norm)
                            st.session_state.speaker_counts = counts_cap10
                            st.session_state.flagged_names = flagged
                        else:
                            if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                                st.session_state.speaker_counts = {}
                            if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                                st.session_state.flagged_names = set()
                    except Exception:
                        pass
                    auto_save()
                    st.rerun()
            else:
                if st.session_state.get("content_type", "Book") == "Script":
                    dialogue_list = extract_dialogue_from_docx_script(st.session_state.docx_path)
                else:
                    dialogue_list = extract_dialogue_from_docx(st.session_state.book_name, st.session_state.docx_path)
                st.session_state.quotes_records = build_quotes_records_from_dialogue_list(dialogue_list)
                st.session_state.quotes_lines = build_quotes_lines_from_records(st.session_state.quotes_records)
                st.session_state.docx_only = True
                st.success("Quotes extracted from DOCX.")
                quotes_txt = "\n".join(dialogue_list)
                st.download_button("Download Extracted Quotes TXT", quotes_txt.encode("utf-8"),
                                   file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt", mime="text/plain")
                if st.button("Restart", key="restart_docx"):
                    restart_app()
                if st.button("Continue", key="continue_docx"):
                    st.session_state.docx_only = False
                    st.session_state.unknown_index = 0
                    st.session_state.console_log = []
                    if st.session_state.get("content_type", "Book") == "Script":
                        st.session_state.step = 3
                    else:
                        st.session_state.step = 2
                    # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
                    try:
                        if st.session_state.get("quotes_lines"):
                            pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                            counts_cap10 = {}
                            flagged = set()
                            for _line in st.session_state.quotes_lines:
                                m = pattern_speaker.match(_line.strip())
                                if not m:
                                    continue
                                speaker_raw = m.group(1).strip()
                                effective = smart_title(speaker_raw)
                                norm = normalize_speaker_name(effective)
                                if norm in flagged:
                                    continue
                                c = counts_cap10.get(norm, 0)
                                if c < 10:
                                    c += 1
                                    counts_cap10[norm] = c
                                    if c >= 10:
                                        flagged.add(norm)
                            st.session_state.speaker_counts = counts_cap10
                            st.session_state.flagged_names = flagged
                        else:
                            if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                                st.session_state.speaker_counts = {}
                            if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                                st.session_state.flagged_names = set()
                    except Exception:
                        pass
                    auto_save()
                    st.rerun()
        else:
            pass
    else:
        docx_file = st.file_uploader("Upload DOCX File", type=["docx"])
        quotes_file = st.file_uploader("Upload Quotes TXT File (optional)", type=["txt"])
        speaker_colors_file = st.file_uploader("Upload Speaker Colors JSON (optional)", type=["json"])
        
        if st.button("Start Processing"):
            if docx_file is None:
                st.error("Please upload a DOCX file.")
            else:
                st.session_state.book_name = os.path.splitext(docx_file.name)[0]
                st.session_state.docx_bytes = docx_file.getvalue()
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                    tmp_docx.write(st.session_state.docx_bytes)
                    st.session_state.docx_path = tmp_docx.name
                if speaker_colors_file is not None:
                    raw = json.load(speaker_colors_file)
                    st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in raw.items()}
                    save_speaker_colors(st.session_state.existing_speaker_colors)
                    st.session_state.speaker_colors = st.session_state.existing_speaker_colors.copy()
                else:
                    st.session_state.existing_speaker_colors = {}
                    st.session_state.speaker_colors = {}
                if quotes_file is not None:
                    quotes_text = quotes_file.read().decode("utf-8")
                    st.session_state.quotes_lines = quotes_text.splitlines(keepends=True)
                    st.session_state.quotes_records = build_quotes_records_from_quotes_lines(st.session_state.quotes_lines)
                    sync_quotes_lines_from_records()
                    st.session_state.docx_only = False
                    # Persist uploaded quotes to a consistent filename and ensure JSON cache for previews
                    try:
                        # Save uploaded quotes to working directory with userkey-bookname naming
                        if 'userkey' in st.session_state and st.session_state.get('book_name'):
                            quotes_filename = f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt"
                        else:
                            # Fallback name if userkey/book_name not set for any reason
                            quotes_filename = "uploaded-quotes.txt"
                        with open(quotes_filename, "w", encoding="utf-8") as _qf:
                            _qf.write(quotes_text)
                        # Ensure the paragraph JSON exists (no DOCX fallback in Step 2)
                        write_paragraph_json_for_session()
                    except Exception as _e:
                        # Do not fail hard; Step 2 will show 'JSON cache not found yet' if this fails
                        pass
                else:
                    st.session_state.quotes_lines = None
                    st.session_state.quotes_records = []
                    st.session_state.docx_only = True
                
                    # Create/overwrite the paragraph JSON once here for docx-only case
                    write_paragraph_json_for_session()
                st.session_state.unknown_index = 0
                st.session_state.console_log = []
                if st.session_state.docx_only:
                    st.session_state.step = 1
                else:
                    if st.session_state.get("content_type", "Book") == "Script":
                        st.session_state.step = 3
                    else:
                        st.session_state.step = 2
                    # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
                    try:
                        if st.session_state.get("quotes_lines"):
                            pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                            counts_cap10 = {}
                            flagged = set()
                            for _line in st.session_state.quotes_lines:
                                m = pattern_speaker.match(_line.strip())
                                if not m:
                                    continue
                                speaker_raw = m.group(1).strip()
                                effective = smart_title(speaker_raw)
                                norm = normalize_speaker_name(effective)
                                if norm in flagged:
                                    continue
                                c = counts_cap10.get(norm, 0)
                                if c < 10:
                                    c += 1
                                    counts_cap10[norm] = c
                                    if c >= 10:
                                        flagged.add(norm)
                            st.session_state.speaker_counts = counts_cap10
                            st.session_state.flagged_names = flagged
                        else:
                            if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                                st.session_state.speaker_counts = {}
                            if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                                st.session_state.flagged_names = set()
                    except Exception:
                        pass
                auto_save()
                st.rerun()

# ========= STEP 2: Quote Record Review =========
elif st.session_state.step == 2:
    ensure_quotes_records_in_session()
    st.markdown("<h4>Step 2: Review Quote Records</h4>", unsafe_allow_html=True)
    st.write("Review each unresolved quote record. Type a speaker, or use 'skip', 'exit', or 'undo'.")
    if st.button("Auto-populate context for all quote records"):
        summary = autopopulate_context_for_all_records(st.session_state.get("quotes_records") or [])
        sync_quotes_lines_from_records()
        auto_save()
        st.success(
            f"Context pass complete: updated {summary['updated']} records "
            f"({summary['with_context']} with context, {summary['without_context']} without context)."
        )

    review_index, review_record = get_next_record_for_review(
        st.session_state.get("quotes_records") or [],
        st.session_state.get("unknown_index", 0),
    )
    if review_index is None or review_record is None:
        st.write("No more unresolved quote records found.")
        if st.button("Proceed to Color Assignment"):
            st.session_state.step = 3
            auto_save()
            st.rerun()
    else:
        dialogue = (review_record.get("quote_with_marks") or review_record.get("quote_text") or "").strip()
        st.markdown("<hr style='margin: 2px 0;'>", unsafe_allow_html=True)
        # Using global JSON-only context resolver
        
        # Compute occurrence target from all previous lines (quoted-segment aware)
        occurrence_target = 1
        start_paragraph_index = 0
        try:
            qrecs = st.session_state.get("quotes_records") or []
            occurrence_target = compute_occurrence_target_for_review(qrecs, review_index)
            start_paragraph_index = compute_start_paragraph_index_for_review(qrecs, review_index)
#            st.session_state._dbg_occurrence_target = occurrence_target
        except Exception:
            pass
        context = get_context_for_dialogue_json_only(
            dialogue,
            occurrence_target=occurrence_target,
            start_paragraph_index=start_paragraph_index,
        )
        if context:
            if "previous" in context:
                st.markdown(neutralize_markdown_in_html(context["previous"]), unsafe_allow_html=True)
            st.markdown(neutralize_markdown_in_html(context["current"]), unsafe_allow_html=True)
#            try:
#                st.text('DEBUG CMP: curr=' + str(st.session_state.get('_dbg_curr_norm'))
#                         + ' | prev1=' + str(st.session_state.get('_dbg_prev1_norm'))
#                         + ' | prev2=' + str(st.session_state.get('_dbg_prev2_norm')))
#            except Exception:
#                pass

            if "next" in context:
                st.markdown(neutralize_markdown_in_html(context["next"]), unsafe_allow_html=True)
        else:
            st.write("No context found in cached JSON for this quote.")

        populate_record_context_fields(review_record, context, occurrence_target)

        st.markdown("<hr style='margin: 2px 0;'>", unsafe_allow_html=True)
        st.write(f"**Dialogue (Line {review_record.get('index', review_index+1)}):** {dialogue}")
        
        def process_unknown_input(new_speaker: str):
            new_speaker = new_speaker.strip()
            if not new_speaker:
                st.session_state.console_log.insert(0, "Empty input ignored. Enter a name, or use 'skip' / 'exit'.")
                return
            if new_speaker.lower() == "exit":
                st.session_state.console_log.insert(0, "Exiting unknown speaker processing.")
                st.session_state.step = 3
            elif new_speaker.lower() == "skip":
                prev_speaker, new_speaker_value = update_record_speaker(review_record, "", action_type="skip")
                append_review_event(review_record, "skip", prev_speaker, new_speaker_value)
                st.session_state.console_log.insert(0, f"Skipped line {review_index+1}.")
                st.session_state.unknown_index = review_index + 1
            elif new_speaker.lower() == "undo":
                if "last_update" in st.session_state and st.session_state.last_update:
                    last = st.session_state.last_update
                    last_index = last.get("index")
                    if last_index is not None and 0 <= last_index < len(st.session_state.get("quotes_records") or []):
                        st.session_state.quotes_records[last_index].update(last.get("record_snapshot") or {})
                        st.session_state.unknown_index = last_index
                        append_review_event(
                            st.session_state.quotes_records[last_index],
                            "undo",
                            last.get("new_speaker"),
                            last.get("previous_speaker"),
                        )
                        st.session_state.console_log.insert(0, f"Reverted line {last_index+1} to previous speaker.")
                    del st.session_state.last_update
                else:
                    st.session_state.console_log.insert(0, "Nothing to undo.")
            else:
                previous_speaker = review_record.get("speaker_text")
                st.session_state.last_update = {
                    "index": review_index,
                    "record_snapshot": dict(review_record),
                    "previous_speaker": previous_speaker,
                    "new_speaker": smart_title(new_speaker),
                }
                updated_speaker = smart_title(new_speaker)

                # Increment count for unflagged speakers and flag at 10
                try:
                    norm = normalize_speaker_name(updated_speaker)
                    if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                        st.session_state.speaker_counts = {}
                    if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                        st.session_state.flagged_names = set()
                    if norm not in st.session_state.flagged_names:
                        new_cnt = st.session_state.speaker_counts.get(norm, 0) + 1
                        if new_cnt >= 10:
                            new_cnt = 10
                            st.session_state.flagged_names.add(norm)
                        st.session_state.speaker_counts[norm] = new_cnt
                except Exception as _e:
                    pass
                prev_speaker, new_speaker_value = update_record_speaker(review_record, updated_speaker, action_type="correct")
                append_review_event(review_record, "correct", prev_speaker, new_speaker_value)
                st.session_state.console_log.insert(0, f"Updated line {review_index+1} with speaker: {updated_speaker}")
                st.session_state.unknown_index = review_index + 1
            sync_quotes_lines_from_records()
            auto_save()
            st.rerun()
        
        # --- New: one‑submit‑per‑name form -------------------------------

        # Frequent speakers (flagged, alphabetical). Buttons act like typing + Enter.
        try:
            if "flagged_names" in st.session_state and st.session_state.flagged_names:
                flagged_sorted = sorted(st.session_state.flagged_names)
                flagged_sorted = [n for n in flagged_sorted if n.lower() != "unknown"]
                st.caption("Frequent speakers:")
                with st.container(horizontal=True):
                    cmap = st.session_state.get("canonical_map") or {}
                    for i, norm in enumerate(flagged_sorted):
                        display_name = cmap.get(norm, norm.title())
                        if st.button(display_name, key=f"flagged_{norm}"):
                            process_unknown_input(display_name)
        except Exception as _e:
            pass
        with st.form("unknown_form", clear_on_submit=True):
            new_name = st.text_input(
                "Enter speaker name (or 'skip'/'exit'/'undo'):",
                key="new_speaker_input",
                placeholder="Type name and press Enter",
            )
        
            # Evenly spaced horizontal buttons
            with st.container(horizontal=True):
                submitted    = st.form_submit_button("Submit")
                skip_clicked = st.form_submit_button("Skip")
                exit_clicked = st.form_submit_button("Exit")
                undo_clicked = st.form_submit_button("Undo (max 1)")
        if submitted:
            process_unknown_input(new_name)
        elif skip_clicked:
            process_unknown_input("skip")
        elif exit_clicked:
            process_unknown_input("exit")
        elif undo_clicked:
            process_unknown_input("undo")
          
        st.text_area("Console Log", "\n".join(st.session_state.console_log), height=150, label_visibility="collapsed")

# ========= STEP 3: Speaker Color Assignment =========
elif st.session_state.step == 3:
    ensure_quotes_records_in_session()
    st.markdown("<h4>Step 3: Speaker Color Assignment</h4>", unsafe_allow_html=True)
    st.write("Assign highlight colors for speakers that do not yet have an assigned color. You can also click 'Edit Speaker Colors' to review and change all assignments.")
    # Load the canonical speakers.
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w+", encoding="utf-8") as tmp_quotes:
        tmp_quotes.write("".join(st.session_state.quotes_lines))
        tmp_quotes_path = tmp_quotes.name
    canonical_speakers, canonical_map = get_canonical_speakers(tmp_quotes_path)
    st.session_state.canonical_map = canonical_map
    # Load existing colors (or default to empty dict)
    existing_colors = st.session_state.get("existing_speaker_colors") or load_existing_colors() or {}
    
    # Determine which speakers need a new assignment
    speakers_to_assign = [
        sp for sp in canonical_speakers 
        if sp.lower() not in ("unknown", "do not read") and (normalize_speaker_name(sp) not in existing_colors or existing_colors.get(normalize_speaker_name(sp), "none") == "none")
    ]
    
    if speakers_to_assign:
        st.write("Assign colors to the following speakers:")
        color_options = [color.title() for color in COLOR_PALETTE.keys() if color.lower() != "do not read"]
        updated_colors = {}
        for sp in speakers_to_assign:
            norm = normalize_speaker_name(sp)
            default_color = existing_colors.get(norm, "none")
            try:
                default_index = color_options.index(default_color.title())
            except ValueError:
                default_index   = color_options.index("None")
            selected = st.selectbox(sp, options=color_options, index=default_index, key="new_"+norm)
            updated_colors[norm] = selected.lower()
        # Merge updated colors with any already assigned values.
        for norm, col in updated_colors.items():
            existing_colors[norm] = col
        # Build final dictionary using normalized keys.
        final_colors = {}
        for sp in canonical_speakers:
            norm = normalize_speaker_name(sp)
            if sp.lower() == "unknown":
                final_colors[norm] = "none"
            elif sp.lower() == "do not read":
                final_colors[norm] = "do not read"
            else:
                final_colors[norm] = existing_colors.get(norm, "none")
        st.session_state.speaker_colors = final_colors
        st.session_state.existing_speaker_colors = existing_colors.copy()
        save_speaker_colors(final_colors)
        st.success("Speaker colors updated.")
    else:
        st.write("All speakers already have assigned colors.")
    if os.path.exists(get_saved_colors_file()):
        with open(get_saved_colors_file(), "r", encoding="utf-8") as f:
            loaded_colors = json.load(f)
        st.session_state.speaker_colors = loaded_colors
        st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in loaded_colors.items()}
    else:
        st.session_state.speaker_colors = {}
        st.session_state.existing_speaker_colors = {}
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Continue"):
            st.session_state.step = 4
            auto_save()
            st.rerun()
    with col2:
        if st.button("Edit Speaker Colors"):
            if os.path.exists(get_saved_colors_file()):
                with open(get_saved_colors_file(), "r", encoding="utf-8") as f:
                    loaded_colors = json.load(f)
                st.session_state.speaker_colors = loaded_colors
                st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in loaded_colors.items()}
            st.session_state.step = "edit_colors"
            auto_save()
            st.rerun()

# ========= EDIT COLORS: Full Speaker Color Assignment =========
elif st.session_state.step == "edit_colors":
    ensure_quotes_records_in_session()
    st.markdown("<h4>Edit Speaker Colors</h4>", unsafe_allow_html=True)
    st.write("Edit the assigned colors for all speakers (excluding 'Unknown'):")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w+", encoding="utf-8") as tmp_quotes:
        tmp_quotes.write("".join(st.session_state.quotes_lines))
        tmp_quotes_path = tmp_quotes.name
    canonical_speakers, canonical_map = get_canonical_speakers(tmp_quotes_path)
    st.session_state.canonical_map = canonical_map
    # Load current colors (or default to empty)
    existing_colors = st.session_state.get("speaker_colors") or load_existing_colors() or {}
    updated_colors = existing_colors.copy()
    color_options = [color.title() for color in COLOR_PALETTE.keys() if color.lower() != "do not read"]
    for sp in canonical_speakers:
        if sp.lower() in ("unknown", "do not read"):
            continue
        norm = normalize_speaker_name(sp)
        default_color = existing_colors.get(norm, "none")
        try:
            default_index = color_options.index(default_color.title())
        except ValueError:
            default_index = color_options.index("None")
        selected = st.selectbox(sp, options=color_options, index=default_index, key="edit_"+norm)
        updated_colors[norm] = selected.lower()
    st.session_state.speaker_colors = updated_colors
    st.session_state.existing_speaker_colors = updated_colors.copy()
    save_speaker_colors(updated_colors)
    st.success("Speaker colors updated.")
    if st.button("Continue"):
        st.session_state.step = 4
        auto_save()
        st.rerun()

# ========= STEP 4: Final HTML Generation =========
elif st.session_state.step == 4:
    ensure_quotes_records_in_session()
    if "speaker_colors" not in st.session_state or st.session_state.speaker_colors is None:
        st.session_state.speaker_colors = load_existing_colors() or {}
    st.markdown("<h4>Step 4: Final HTML Generation</h4>", unsafe_allow_html=True)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w+", encoding="utf-8") as tmp_quotes:
        tmp_quotes.write("".join(st.session_state.quotes_lines))
        quotes_file_path = tmp_quotes.name
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_marker:
        marker_docx_path = tmp_marker.name
    create_marker_docx(st.session_state.docx_path, marker_docx_path)
    html = convert_docx_to_html_mammoth(marker_docx_path)
    os.remove(marker_docx_path)
    quotes_list = load_quotes(quotes_file_path, st.session_state.canonical_map)
    highlighted_html = highlight_dialogue_in_html(html, quotes_list, st.session_state.speaker_colors)
    final_html_body = apply_manual_indentation_with_markers(st.session_state.docx_path, highlighted_html)
    if st.session_state.get("content_type", "Book") == "Script":
        final_html_body = transform_script_layout(final_html_body)
    summary_html = generate_summary_html(quotes_list, list(st.session_state.canonical_map.values()), st.session_state.speaker_colors)
    ranking_html = generate_ranking_html(quotes_list, st.session_state.speaker_colors)
    first_lines_html = generate_first_lines_html(quotes_list, list(st.session_state.canonical_map.values()))
    fontsel = normalize_font_family(st.session_state.get("fontsel", "Avenir"))
    font_face_html = build_font_face_css(fontsel, embed_base64=True)
    final_html_body = summary_html + "\n<br><br><br>\n" + ranking_html + "\n<br><br><br>\n" + first_lines_html + "\n" + final_html_body
    final_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{st.session_state.book_name}</title>
  <style>
    {font_face_html}
    body {{
      font-family: '{fontsel}', sans-serif;
      line-height: 2;
      max-width: 500px;
      margin: auto;
    }}
    span {{
      padding: 0;
    }}
    span.highlight {{
      background-color: var(--highlight-color, transparent);
      padding: 0.33em 0px;
      box-decoration-break: clone;
      -webkit-box-decoration-break: clone;
    }}

    /* Script layout */
    p.script-line {{
      margin-left: 0;
      text-indent: 0;
      display: grid;
      grid-template-columns: 9em 1fr;  /* fixed speaker column width */
      column-gap: 0.75em;
    }}
    .script-speaker {{
      font-weight: bold;
    }}
    .script-dialogue {{
      /* dialogue automatically takes remaining space in the second column */
    }}

  </style>
</head>
<body>
{final_html_body}
</body>
</html>
"""
    final_html_path = os.path.join(tempfile.gettempdir(), f"{st.session_state.book_name}.html")
    with open(final_html_path, "w", encoding="utf-8") as f:
        f.write(final_html)
    st.success("Final HTML generated.")
    components.html(final_html, height=800, scrolling=True)
    with open(final_html_path, "rb") as f:
        html_bytes = f.read()
    st.download_button("Download HTML File", html_bytes,
                       file_name=f"{st.session_state.userkey}-{st.session_state.book_name}.html", mime="text/html")
    # --- PDF export (optional) ---
    pdf_file_name = f"{st.session_state.userkey}-{st.session_state.book_name}.pdf"

    # Only *check* availability during Step 4 render (fast). Actual PDF generation happens on-click.
    pdf_available = True
    pdf_import_error: str | None = None
    try:
        import weasyprint  # type: ignore  # noqa: F401
    except Exception as e:
        pdf_available = False
        pdf_import_error = str(e)

    if pdf_available:
        def _make_pdf() -> bytes:
            # Streamlit can lazily call this when the download button is clicked (newer versions).
            # For older Streamlit versions (no callable support), we fall back below.
            return render_html_to_pdf_bytes(final_html, base_url=os.path.dirname(final_html_path))

        try:
            st.download_button(
                "Download PDF File (takes a while!)",
                data=_make_pdf,  # lazy / on-click generation (Streamlit >= supports callable)
                file_name=pdf_file_name,
                mime="application/pdf",
            )
        except TypeError:
            # Fallback for older Streamlit: generate eagerly (but still only if this branch is reached).
            try:
                pdf_bytes = _make_pdf()
                st.download_button(
                    "Download PDF File (takes a while!)",
                    pdf_bytes,
                    file_name=pdf_file_name,
                    mime="application/pdf",
                )
            except Exception as e:
                st.download_button(
                    "Download PDF File (takes a while!)",
                    b"",
                    file_name=pdf_file_name,
                    mime="application/pdf",
                    disabled=True,
                )
                st.caption(f"PDF export failed: {e}")
    else:
        st.download_button(
            "Download PDF File (takes a while!)",
            b"",
            file_name=pdf_file_name,
            mime="application/pdf",
            disabled=True,
        )
        st.caption(
            "PDF export is unavailable in this environment. "
            "Install WeasyPrint (plus its system dependencies) to enable it. "
            f"Details: {pdf_import_error}"
        )

    updated_colors = json.dumps(st.session_state.speaker_colors, indent=4, ensure_ascii=False).encode("utf-8")
    st.download_button("Download Updated Speaker Colors JSON", updated_colors,
                       file_name=f"{st.session_state.userkey}-speaker_colors.json", mime="application/json")
    updated_quotes = "".join(st.session_state.quotes_lines).encode("utf-8")
    st.download_button("Download Updated Quotes TXT", updated_quotes,
                       file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt", mime="text/plain")
    quotes_records_payload = st.session_state.get("quotes_records") or []
    quotes_records_bytes = json.dumps(quotes_records_payload, indent=2, ensure_ascii=False).encode("utf-8")
    st.download_button(
        "Download Quotes Records JSON",
        quotes_records_bytes,
        file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes-records.json",
        mime="application/json",
    )
    # Lines CSV export (eager generation; lightweight and avoids Streamlit context issues in deferred callables)
    csv_bytes = build_csv_from_docx_json_and_quotes()
    st.download_button(
        "Download Lines CSV",
        csv_bytes,
        file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-lines.csv",
        mime="text/csv",
    )
    if os.path.exists(get_unmatched_quotes_filename()):
        with open(get_unmatched_quotes_filename(), "rb") as f:
            unmatched_bytes = f.read()
        st.download_button("Download Unmatched Quotes TXT", unmatched_bytes,
                           file_name=get_unmatched_quotes_filename(), mime="text/plain")
    if st.button("Return to Step 2"):
        if "book_name" in st.session_state:
            quotes_filename = f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt"
            if os.path.exists(quotes_filename):
                with open(quotes_filename, "r", encoding="utf-8") as f:
                    st.session_state.quotes_lines = f.read().splitlines(keepends=True)
                st.session_state.quotes_records = build_quotes_records_from_quotes_lines(st.session_state.quotes_lines)
                sync_quotes_lines_from_records()
        if os.path.exists(f"{st.session_state.userkey}-speaker_colors.json"):
            with open(f"{st.session_state.userkey}-speaker_colors.json", "r", encoding="utf-8") as f:
                colors = json.load(f)
            st.session_state.speaker_colors = colors
            st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in colors.items()}
        st.session_state.step = 2
        # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
        try:
            if st.session_state.get("quotes_lines"):
                pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                counts_cap10 = {}
                flagged = set()
                for _line in st.session_state.quotes_lines:
                    m = pattern_speaker.match(_line.strip())
                    if not m:
                        continue
                    speaker_raw = m.group(1).strip()
                    effective = smart_title(speaker_raw)
                    norm = normalize_speaker_name(effective)
                    if norm in flagged:
                        continue
                    c = counts_cap10.get(norm, 0)
                    if c < 10:
                        c += 1
                        counts_cap10[norm] = c
                        if c >= 10:
                            flagged.add(norm)
                st.session_state.speaker_counts = counts_cap10
                st.session_state.flagged_names = flagged
            else:
                if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                    st.session_state.speaker_counts = {}
                if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                    st.session_state.flagged_names = set()
        except Exception:
            pass
        auto_save()
        st.rerun() 
        # Add a Clear Cache button below "Return to Step 2"
    if st.button("Clear Cache for This User"):
        # Only delete files for this userkey
        files_to_remove = [
            get_progress_file(),
            get_saved_colors_file(),
            get_unmatched_quotes_filename(),
            f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt",
            f"{st.session_state.userkey}-{st.session_state.book_name}.html",
            f"{st.session_state.userkey}-{st.session_state.book_name}.json"
        ]
        for path in files_to_remove:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception as e:
                st.warning(f"Could not remove {path}: {e}")

        # List all your app’s keys here to reset ONLY relevant state
        keys_to_clear = [
            "step", "userkey", "docx_bytes", "docx_path", "book_name",
            "quotes_lines", "speaker_colors", "existing_speaker_colors",
            "unknown_index", "console_log", "canonical_map", "last_update"
        ]
        for k in keys_to_clear:
            if k in st.session_state:
                del st.session_state[k]
        st.rerun()
