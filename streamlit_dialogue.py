import streamlit as st
import re

def extract_italic_spans(para):
    """Return list of ((start, end), text) for italic content in this paragraph,
    with spans measured against para.text.
    We consider a run italic if any of run.italic, run.font.italic is True,
    or the character style name contains 'Italic'/'Emphasis' (best-effort)."""
    runs = getattr(para, 'runs', [])
    spans = []
    pos = 0
    # Build a mirror of para.text by concatenating run.text, tracking offsets
    for run in runs:
        t = run.text or ''
        length = len(t)
        italic_flag = False
        try:
            if run.italic is True:
                italic_flag = True
        except Exception:
            pass
        try:
            if getattr(run.font, 'italic', None) is True:
                italic_flag = True
        except Exception:
            pass
        try:
            sty = getattr(run, 'style', None)
            if sty is not None:
                name = getattr(sty, 'name', '') or ''
                if 'italic' in name.lower() or 'emphasis' in name.lower():
                    italic_flag = True
        except Exception:
            pass
        if italic_flag and length > 0:
            spans.append(((pos, pos + length), t))
        pos += length
    # Merge adjacent/contiguous italic runs
    merged = []
    for span, seg in spans:
        if not merged:
            merged.append([list(span), seg])
        else:
            last_span, last_text = merged[-1]
            if span[0] == last_span[1]:
                last_span[1] = span[1]
                merged[-1][1] = last_text + seg
            else:
                merged.append([list(span), seg])
    return [((s[0], s[1]), txt) for s, txt in merged]

import os
import json
import tempfile
import mammoth
import docx
import base64
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup, NavigableString
from collections import Counter
import html

def trim_paragraph_cache_before_previous(previous_html: str):
    try:
        djson_path = st.session_state.get('d_json_path')
        if not djson_path or not os.path.exists(djson_path):
            return False
        with open(djson_path, "r", encoding="utf-8") as f:
            paragraphs_html = json.load(f)
        if not previous_html:
            return False

        # Exact match first
        try:
            idx = paragraphs_html.index(previous_html)
        except ValueError:
            # Fallback: compare by text (strip tags) to be resilient
            def strip_html(s):
                try:
                    soup = BeautifulSoup(s, "html.parser")
                    return soup.get_text() or ""
                except Exception:
                    return s
            prev_text = strip_html(previous_html)
            idx = -1
            for i, p in enumerate(paragraphs_html):
                if strip_html(p) == prev_text:
                    idx = i
                    break

        if idx <= 0:
            # Either not found (-1) or already first item (0); nothing to remove before it.
            return False

        # Keep from 'previous_html' onward (i.e., drop 0..idx-1 inclusive)
        new_paragraphs = paragraphs_html[idx:]
        with open(djson_path, "w", encoding="utf-8") as f:
            json.dump(new_paragraphs, f, ensure_ascii=False)
        # Update session so subsequent reads are consistent
        st.session_state.trimmed_paragraphs_since_last = True
        return True
    except Exception:
        return False

def neutralize_markdown_in_html(html_s: str) -> str:
    try:
        soup = BeautifulSoup(html_s, "html.parser")
        for t in list(soup.find_all(string=True)):
            # Skip script/style tags
            if getattr(t.parent, "name", "").lower() in ("script", "style"):
                continue
            # Replace literal '*' and '_' with HTML entities so st.markdown doesn't italicise them.
            new_txt = t.replace("*", r"\*").replace("_", r"\_")
            if new_txt != t:
                t.replace_with(new_txt)
        return str(soup)
    except Exception:
        # Fallback: raw replacements (may affect tags if present, but better than italics)
        return html_s.replace("*", "&#42;").replace("_", "&#95;")

# Import components for HTML embedding.
import streamlit.components.v1 as components

def build_d_paragraphs_html(docx_path):
    import docx
    import html
    try:
        doc = docx.Document(docx_path)
    except Exception:
        return []
    def wrap(txt, b, i, u):
        s = html.escape(txt or "")
        if not s:
            return s
        if u: s = f"<u>{s}</u>"
        if i: s = f"<i>{s}</i>"
        if b: s = f"<b>{s}</b>"
        return s
    out = []
    # ensure `re` is imported at top of file: import re
    for p in doc.paragraphs:
        if not p.runs:
            candidate = html.escape(p.text or "")
        else:
            candidate = "".join(
                wrap(r.text, r.bold, r.italic, r.underline) for r in p.runs
            )
    
        # Drop empty/whitespace-only paragraphs (ignoring any HTML tags)
        plain = re.sub(r"<[^>]*>", "", candidate)  # strip tags for the emptiness check
        if not re.search(r"\S", plain):            # no non-whitespace char
            continue
    
        out.append(candidate)
    return out



def get_context_for_dialogue_json_only(dialogue: str, occurrence_target: int = 1):
    try:
        djson_path = st.session_state.get('d_json_path')
        if not djson_path or not os.path.exists(djson_path):
            return None
        with open(djson_path, "r", encoding="utf-8") as f:
            paragraphs_html = json.load(f)  # list[str]
    except Exception:
        return None

    def soup_text(html_s: str) -> str:
        try:
            soup = BeautifulSoup(html_s, "html.parser")
            return soup.get_text() or ""
        except Exception:
            return html_s

    dlg = dialogue
    m_q = re.search(r'[“"]([^”"]+)[”"]', dlg) or re.search(r"[‘']([^’']+)[’']", dlg)
    dialogue_to_highlight = m_q.group(1) if m_q else dlg
    normalized_highlight = normalize_text(dialogue_to_highlight).lower()

    if not normalized_highlight.strip():
        occurrence_target = 1

    plain_paras = [soup_text(p) for p in paragraphs_html]

    cumulative = 0
    chosen_idx = None
    within_para_target = 1

    for idx, para_plain in enumerate(plain_paras):
        para_norm = normalize_text(para_plain).lower()
        count_here = len(re.findall(re.escape(normalized_highlight), para_norm)) if normalized_highlight else 0
        if count_here > 0:
            if cumulative + count_here >= occurrence_target:
                chosen_idx = idx
                within_para_target = occurrence_target - cumulative
                break
            cumulative += count_here

    if chosen_idx is None:
        for idx, para_plain in enumerate(plain_paras):
            if normalized_highlight in normalize_text(para_plain).lower():
                chosen_idx = idx
                within_para_target = 1
                break

    if chosen_idx is None:
        return None

    ctx = {}
    if chosen_idx > 0:
        ctx["previous"] = paragraphs_html[chosen_idx - 1]

    try:
        soup = BeautifulSoup(paragraphs_html[chosen_idx], "html.parser")
                # --- begin: tag-stripped match + bold in original HTML (early-return if applied) ---
        # 1) Build a tag-stripped view for matching (keeps text as user sees it)
        plain_text = soup.get_text()
        
        # 2) Find the m-th (within_para_target) occurrence on the stripped view
        _pat = re.compile(re.escape(dialogue_to_highlight), re.IGNORECASE)
        _occ = 0
        _span = None
        for _m in _pat.finditer(plain_text):
            _occ += 1
            if _occ == within_para_target:
                _span = (_m.start(), _m.end())
                break
        
        # 3) If found, map the [start:end) range back onto the original soup's text nodes and wrap with <b>
        if _span is not None:
            _start, _end = _span
            running = 0
            # Walk text nodes in order; split and wrap overlap segments with <b>
            for tnode in list(soup.find_all(string=True)):
                p = getattr(tnode.parent, "name", "").lower()
                if p in ("script", "style"):
                    continue
                text = str(tnode)
                length = len(text)
                node_start = running
                node_end = running + length
        
                # Does this node overlap the target [start, end)?
                if node_end > _start and node_start < _end:
                    # overlap in this node
                    ov_start = max(_start, node_start)
                    ov_end   = min(_end, node_end)
                    rel_start = ov_start - node_start
                    rel_end   = ov_end   - node_start
        
                    before = text[:rel_start]
                    match  = text[rel_start:rel_end]
                    after  = text[rel_end:]
        
                    from bs4 import NavigableString
                    new_nodes = []
                    if before:
                        new_nodes.append(NavigableString(before))
                    btag = soup.new_tag("b")
                    btag.string = match
                    new_nodes.append(btag)
                    if after:
                        new_nodes.append(NavigableString(after))
        
                    # Replace this text node with the split nodes
                    tnode.replace_with(*new_nodes)
        
                    # IMPORTANT: if match spans multiple consecutive text nodes,
                    # decrease the remaining range and continue across nodes
                    # Update the remaining global range for subsequent nodes:
                    # we wrapped [ov_start, ov_end), so advance _start to ov_end
                    _start = ov_end
                    if _start >= _end:
                        break  # fully wrapped
        
                    # Adjust running to reflect we've consumed this node
                    running = node_end
                    continue
        
                running += length
        
            ctx["current"] = str(soup)
            if chosen_idx + 1 < len(paragraphs_html):
                ctx["next"] = paragraphs_html[chosen_idx + 1]
            return ctx
        # --- end: tag-stripped match + bold in original HTML ---
        pattern = re.compile(re.escape(dialogue_to_highlight), re.IGNORECASE)

        global_counter = 0
        replaced_flag = False

        def replace_only_mth(node):
            nonlocal global_counter, replaced_flag
            text_val = str(node)
            parts = []
            last = 0
            for m in pattern.finditer(text_val):
                global_counter += 1
                parts.append(text_val[last:m.start()])
                seg = m.group(0)
                if (not replaced_flag) and global_counter == within_para_target:
                    b = soup.new_tag("b")
                    b.string = seg
                    parts.append(b)
                    replaced_flag = True
                else:
                    parts.append(seg)
                last = m.end()
            parts.append(text_val[last:])
            if len(parts) > 1:
                anchor = None
                for i, part in enumerate(parts):
                    new_node = soup.new_string(part) if isinstance(part, str) else part
                    if i == 0:
                        node.replace_with(new_node)
                        anchor = new_node
                    else:
                        anchor.insert_after(new_node)
                        anchor = new_node

        for tnode in list(soup.find_all(string=True)):
            parent_name = getattr(tnode.parent, "name", "").lower()
            if parent_name in ("script", "style"):
                continue
            if replaced_flag:
                continue
            replace_only_mth(tnode)

        if not replaced_flag:
            global_counter = 0
            replaced_flag = False
            for tnode in list(soup.find_all(string=True)):
                parent_name = getattr(tnode.parent, "name", "").lower()
                if parent_name in ("script", "style"):
                    continue
                replace_only_mth(tnode)
                if replaced_flag:
                    break

        ctx["current"] = str(soup)
    except Exception:
        ctx["current"] = paragraphs_html[chosen_idx]

    if chosen_idx + 1 < len(paragraphs_html):
        ctx["next"] = paragraphs_html[chosen_idx + 1]

    return ctx


#def compute_and_set_d_json_path():
#    """Ensure JSON paragraph cache exists and set st.session_state['d_json_path'].
#    JSON is written alongside the DOCX using the book name. No DOCX fallback in Step 2.
#    """
#    try:
#        docx_path = st.session_state.get("docx_path")
#        book_name = st.session_state.get("book_name") or "document"
#        if not docx_path or not os.path.exists(docx_path):
#            return
#       folder = os.path.dirname(docx_path)
#        json_path = os.path.join(folder, f"{book_name}.json")
#        paragraphs = build_d_paragraphs_html(docx_path)
#        with open(json_path, "w", encoding="utf-8") as f:
#            json.dump(paragraphs, f, ensure_ascii=False, indent=2)
#        st.session_state['d_json_path'] = json_path
#    except Exception:
#        # Do not fall back; if JSON generation fails, leave preview blank.
#        pass

def write_paragraph_json_for_session():
    """Create or overwrite [userkey]-[bookname].json in the working directory
    from st.session_state['docx_path'], set st.session_state['d_json_path'],
    and never fall back to DOCX later. This runs only when 'Start Processing' is clicked.
    """
    try:
        import os, json
        docx_path = st.session_state.get('docx_path')
        if not docx_path or not os.path.exists(docx_path):
            return
        userkey = st.session_state.get('userkey') or "User"
        book_name = st.session_state.get('book_name') or "document"
        json_path = os.path.join(os.getcwd(), f"{userkey}-{book_name}.json")
        paragraphs = build_d_paragraphs_html(docx_path)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(paragraphs, f, ensure_ascii=False, indent=2)
        st.session_state['d_json_path'] = json_path
    except Exception:
        # Do not silently recreate elsewhere; leave preview blank if this fails.
        pass

#def ensure_d_json(docx_path, quotes_path):
#    """Deprecated: use write_paragraph_json_for_session(). Keeping for backward compatibility."""
#    write_paragraph_json_for_session()


# Inject custom CSS
custom_css = """
<style>
:root {
  --primary-color: #008080;      /* Teal */
  --primary-hover: #007070;
  --background-color: #fdfdfd;
  --text-color: #222222;
  --card-background: #ffffff;
  --card-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
  --border-radius: 10px;
  --font-family: 'Avenir', sans-serif;
  --accent-color: #ff9900;
}

/* Global Styles */
body {
  background-color: var(--background-color);
  font-family: var(--font-family);
  color: var(--text-color);
  margin: 0;
  padding: 0;
}

h1, h2, h3, h4, h5, h6 {
  color: var(--text-color);
  font-weight: 700;
  margin-bottom: 0.5em;
}

/* Button Styles */
div.stButton > button {
  background-color: var(--primary-color);
  color: #ffffff;
  border: none;
  padding: 0.75em 1.25em;
  border-radius: var(--border-radius);
  cursor: pointer;
  transition: background-color 0.3s ease, transform 0.2s;
}

div.stButton > button:hover {
  background-color: var(--primary-hover);
  transform: translateY(-2px);
}

/* Card/Container Styling */
.custom-container {
  background: var(--card-background);
  padding: 2em;
  border-radius: var(--border-radius);
  box-shadow: var(--card-shadow);
  margin-bottom: 2em;
}

.css-1d391kg {
  background: var(--card-background);
  padding: 1em;
  border-radius: var(--border-radius);
  box-shadow: var(--card-shadow);
}

/* Form Element Styling */
input, select, textarea {
  border: 1px solid #ddd;
  border-radius: 4px;
  padding: 0.5em;
  font-size: 1em;
}

input:focus, select:focus, textarea:focus {
  outline: none;
  border-color: var(--primary-color);
  box-shadow: 0 0 5px rgba(0, 128, 128, 0.3);
}
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ---------------------------
# Global Constants & Helper Functions
# ---------------------------

COLOR_PALETTE = {
    "dark grey": (180, 178, 179, 1, "rgb(30, 28, 29)"),
    "burgundy": (207, 156, 153, 1, "rgb(55, 4, 1)"),
    "red": (255, 153, 153, 1, "rgb(107, 9, 6)"),
    "orange": (255, 198, 153, 1, "rgb(119, 62, 17)"),
    "yellow": (255, 241, 153, 1, "rgb(107, 93, 5)"),
    "dark yellow": (230, 223, 153, 1, "rgb(78, 71, 1)"),
    "brown": (205, 192, 173, 1, "rgb(61, 48, 29)"),
    "silver": (243, 243, 243, 1, "rgb(96, 96, 96)"),
    "light green": (160, 255, 153, 1, "rgb(20, 115, 13)"),
    "dark green": (157, 192, 152, 1, "rgb(7, 42, 2)"),
    "turquoise": (153, 255, 234, 1, "rgb(3, 105, 84)"),
    "light blue": (153, 248, 254, 1, "rgb(9, 104, 110)"),
    "bright blue": (152, 166, 255, 1, "rgb(4, 18, 107)"),
    "dark blue": (152, 166, 255, 1, "rgb(4, 18, 107)"),
    "navy blue": (153, 162, 201, 1, "rgb(4, 13, 52)"),
    "dark purple": (219, 173, 222, 1, "rgb(68, 22, 71)"),
    "light purple": (231, 203, 254, 1, "rgb(81, 53, 104)"),
    "bright pink": (255, 153, 236, 1, "rgb(112, 10, 93)"),
    "light pink": (254, 238, 248, 1, "rgb(103, 87, 97)"),
    "pale pink": (254, 238, 248, 1, "rgb(103, 87, 97)"),
    "wine": (234, 184, 185, 1, "rgb(90, 40, 41)"),
    "lime": (230, 244, 182, 1, "rgb(81, 95, 33)"),
    "none": (134, 8, 0, 1.0, "rgb(134, 8, 0)"),
    "do not read": (0, 0, 0, 1, "rgb(100, 100, 100)"),
    "error": (0, 0, 0, 0, "")  # For "Error": transparent background, no text color override.
}


# ==== STEP 0: Userkey Entry ====
if "userkey" not in st.session_state:
    st.session_state.userkey = ""

if "step" not in st.session_state:
    st.session_state.step = 0

# ========= STEP 0: User Identification =========

if st.session_state.step == 0:
    st.title("Welcome to Scripter")
    st.write("Please enter a unique identifier. This can be any memorable username or passphrase. It must be unique to you—do not share it.")
    user_input = st.text_input("Enter your user key (username, nickname, or passphrase):", key="userkey_input")
    if st.button("Next"):
        if user_input.strip() == "":
            st.warning("You must enter a user key to continue.")
            st.stop()
        else:
            st.session_state.userkey = user_input.strip()
            st.session_state.step = 1
            st.rerun()
    st.stop()
    
def get_saved_colors_file():
    return f"{st.session_state.userkey}-speaker_colors.json"
    
def get_progress_file():
    return f"{st.session_state.userkey}-progress.json"
    
def get_unmatched_quotes_filename():
    return f"{st.session_state.userkey}-unmatched_quotes.txt"

def normalize_text(text):
    text = text.replace("\u00A0", " ")
    text = text.replace("…", "...")
    text = text.replace("“", "\"").replace("”", "\"")
    text = text.replace("’", "'").replace("‘", "'")
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def match_normalize(text):
    return text.replace("’", "'").replace("‘", "'")

def normalize_speaker_name(name):
    # Replace typographic apostrophes with straight ones, remove periods, lowercase, and trim.
    return name.replace("’", "'").replace("‘", "'").replace(".", "").lower().strip()

def smart_title(name):
    words = name.split()
    if not words:
        return name
    exceptions = {"ps", "pc", "ds", "di", "dci"}
    new_words = []
    for w in words:
        if w.lower() in exceptions:
            new_words.append(w.upper())
        else:
            new_words.append(w.capitalize())
    result = " ".join(new_words)
    result = re.sub(r"\(([mf])\)$", lambda m: "(" + m.group(1).upper() + ")", result, flags=re.IGNORECASE)
    return result

#def write_file_atomic(filepath, lines):
#    with open(filepath, "w", encoding="utf-8") as f:
#        f.writelines(lines)
#        f.flush()
#        os.fsync(f.fileno())

# ---------------------------
# Auto-Save & Auto-Load Functions
# ---------------------------
def auto_save():
    data = {
        "step": st.session_state.get("step", 1),
        "quotes_lines": st.session_state.get("quotes_lines"),
        "speaker_colors": st.session_state.get("speaker_colors"),
        "unknown_index": st.session_state.get("unknown_index", 0),
        "console_log": st.session_state.get("console_log", []),
        "canonical_map": st.session_state.get("canonical_map") or {},
        "book_name": st.session_state.get("book_name"),
        "existing_speaker_colors": st.session_state.get("existing_speaker_colors")
    }
    if "docx_bytes" in st.session_state and st.session_state.docx_bytes is not None:
        data["docx_bytes"] = base64.b64encode(st.session_state.docx_bytes).decode("utf-8")
    with open(get_progress_file(), "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)
    if st.session_state.get("speaker_colors") is not None:
        with open(get_saved_colors_file(), "w", encoding="utf-8") as f:
            json.dump(st.session_state.speaker_colors, f, indent=4, ensure_ascii=False)
    if st.session_state.get("quotes_lines") and st.session_state.get("book_name"):
        quotes_filename = f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt"
        with open(quotes_filename, "w", encoding="utf-8") as f:
            quotes_text = "".join(st.session_state.quotes_lines)
            f.write(quotes_text)

def auto_load():
    if os.path.exists(get_progress_file()):
        with open(get_progress_file(), "r", encoding="utf-8") as f:
            data = json.load(f)
        for key, value in data.items():
            st.session_state[key] = value
            # Normalise restored structures
            if isinstance(st.session_state.get("flagged_names"), list):
                st.session_state.flagged_names = set(st.session_state.flagged_names)
            if st.session_state.get("speaker_counts") is None:
                st.session_state.speaker_counts = {}
            if st.session_state.get("flagged_names") is None:
                st.session_state.flagged_names = set()
            if st.session_state.get("canonical_map") is None:
                st.session_state.canonical_map = {}

            # Rebuild counts/flags from quotes_lines if missing or empty
            needs_rebuild = (
                not st.session_state.speaker_counts or
                (not st.session_state.flagged_names and st.session_state.speaker_counts)
            )
            if needs_rebuild and st.session_state.get("quotes_lines"):
                pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                counts_cap10 = {}
                flagged = set()
                for _line in st.session_state.quotes_lines:
                    m = pattern_speaker.match(_line.strip())
                    if not m:
                        continue
                    speaker_raw = m.group(1).strip()
                    effective = smart_title(speaker_raw)
                    norm = normalize_speaker_name(effective)
                    if norm in flagged:
                        continue
                    c = counts_cap10.get(norm, 0)
                    if c < 10:
                        c += 1
                        counts_cap10[norm] = c
                        if c >= 10:
                            flagged.add(norm)
                st.session_state.speaker_counts = counts_cap10
                st.session_state.flagged_names = flagged

        if "existing_speaker_colors" in st.session_state and st.session_state.existing_speaker_colors:
            st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in st.session_state.existing_speaker_colors.items()}
        if "docx_bytes" in st.session_state:
            docx_bytes = base64.b64decode(st.session_state["docx_bytes"].encode("utf-8"))
            st.session_state.docx_bytes = docx_bytes
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                tmp_docx.write(docx_bytes)
                st.session_state.docx_path = tmp_docx.name
            # Ensure d_json_path points to the unified JSON cache: [userkey]-[book_name].json in CWD
            try:
                userkey = st.session_state.get("userkey")
                book_name = st.session_state.get("book_name")
                if userkey and book_name:
                    cand = os.path.join(os.getcwd(), f"{userkey}-{book_name}.json")
                    if os.path.exists(cand):
                        st.session_state['d_json_path'] = cand
                    else:
                        # If the JSON cache is missing but we have a docx, rebuild it once
                        if st.session_state.get('docx_path') and os.path.exists(st.session_state['docx_path']):
                            write_paragraph_json_for_session()
            except Exception:
                pass


if os.path.exists(get_progress_file()):
    if st.button("Load Saved Progress"):
         auto_load()
         st.rerun()

# ---------------------------
# Alternative Extraction Functions
# ---------------------------
ATTACH_NO_SPACE = {"'", "’", "‘", '"', "“", "”", ",", ".", ";", ":", "?", "!"}
DASHES = {"-", "–", "—"}

def smart_join(run_texts):
    if not run_texts:
        return ""
    result = run_texts[0]
    for text in run_texts[1:]:
        if not text:
            continue

        # preserve any explicit/trailing spaces already in the text
        if result and result[-1].isspace():
            result += text.lstrip()
            continue
        if text[0].isspace():
            result += text
            continue

        # ellipses should attach to the previous token
        if text.startswith("...") or text.startswith("…"):
            result = result.rstrip()
            result += text
            continue

        prev = result[-1]
        first = text[0]

        # 1) Contractions/possessives: apostrophe binds to following letters (That’s, you’re, I’ll)
        if prev in {"'", "’", "‘"} and first.isalnum():
            result += text
            continue

        # 2) Characters that attach to the previous token (no leading space)
        if first in ATTACH_NO_SPACE:
            result += text
            continue

        # 3) Dashes/hyphens: attach tightly on both sides
        if first in DASHES:
            result += text            # no space before dash
            continue
        if prev in DASHES:
            result += text            # no space after dash
            continue

        # 4) Default spacing rule
        if prev.isalnum() and first.isalnum():
            result += text            # join words without extra space
        else:
            # Added guards (surgical): avoid space before ™/® and inside parentheses
            prev = result[-1] if result else ''
            first = text[0] if text else ''
            if first in {'™','®'} or prev == '(' or first == ')':
                result += text
                continue
            # Guard: keep opening double-quote tight with the next token
            prev = result[-1] if result else ''
            if prev in {'“', '"'}:
                result += text
                continue
            result += " " + text       # otherwise, insert a space
    return result

# def is_run_italic(run):
#    """Return True if the run is italic due to direct formatting or its character style."""
#    try:
#        if getattr(run.font, "italic", None) is True:
#            return True
#        style = getattr(run, "style", None)
#        if style is not None and getattr(style.font, "italic", None) is True:
#            return True
#    except Exception:
#        # Be conservative; if anything goes wrong, treat as non-italic
#        return False
#    return False
def effective_run_italic(run, paragraph):
    """Return True if, after cascading styles, this run is italic.
    Precedence (lowest to highest): paragraph style -> run character style -> direct run formatting.
    Explicit False overrides inherited True.
    """
    base = False
    try:
        # Paragraph style
        psty = getattr(paragraph, "style", None)
        if psty is not None and getattr(psty.font, "italic", None) is True:
            base = True
        # Character style on run
        rsty = getattr(run, "style", None)
        rsty_italic = getattr(getattr(rsty, "font", None), "italic", None)
        if rsty_italic is not None:
            base = bool(rsty_italic)
        # Direct run formatting
        rfmt_italic = getattr(getattr(run, "font", None), "italic", None)
        if rfmt_italic is not None:
            base = bool(rfmt_italic)
    except Exception:
        pass
    return base


def extract_italicized_text(paragraph):
    """
    def _trim_quote_edges(s: str) -> str:
        # Italics-path parity: remove spaces just inside opening/closing double quotes
        s = re.sub(r'(?<=[“\"])\\s+', '', s)
        s = re.sub(r'\\s+(?=[”\"])', '', s)
        return s
    Return a list of italic blocks for a paragraph.
    Detects italics after cascading: paragraph style -> character style -> direct run formatting.
    Preserves the existing >= 2-word threshold and smart_join behaviour.
    """
    italic_blocks = []
    current_block = []
    for run in paragraph.runs:
        if effective_run_italic(run, paragraph):
            current_block.append(run.text)
        else:
            joined = smart_join(current_block)
            joined = _trim_quote_edges(joined)
            if len(joined.split()) >= 2:
                italic_blocks.append(joined)
            current_block = []
    # flush tail
    joined = smart_join(current_block)
    joined = _trim_quote_edges(joined)
    joined = re.sub(r'^\.\s+(?=\w)', '', joined)
    if len(joined.split()) >= 2:
        italic_blocks.append(joined)
    return italic_blocks
def extract_italic_spans(paragraph):
    """
    Return a list of ((start, end), text) for contiguous italic blocks in this paragraph,
    using the same cascade logic and >=2-word rule as extract_italicized_text.
    Spans are computed against the raw concatenation of run.text (paragraph.text),
    with adjustment if a leading ". " is stripped.
    """
    spans = []
    pos = 0
    block_start = None
    block_raw = []  # raw run.text pieces

    for run in paragraph.runs:
        t = run.text or ""
        n = len(t)
        if effective_run_italic(run, paragraph):
            if block_start is None:
                block_start = pos
            block_raw.append(t)
        else:
            if block_start is not None:
                raw = "".join(block_raw)
                joined = smart_join(block_raw)
                shift = 0
                if re.match(r'^\.\s+(?=\w)', joined):
                    shift = 2
                    joined = joined[2:]
                if len(joined.split()) >= 2:
                    start = block_start + shift
                    end = start + max(0, len(raw) - shift)
                    spans.append(((start, end), joined))
                block_start = None
                block_raw = []
        pos += n

    if block_start is not None:
        raw = "".join(block_raw)
        joined = smart_join(block_raw)
        shift = 0
        if re.match(r'^\.\s+(?=\w)', joined):
            shift = 2
            joined = joined[2:]
        if len(joined.split()) >= 2:
            start = block_start + shift
            end = start + max(0, len(raw) - shift)
            spans.append(((start, end), joined))

    return spans
def extract_dialogue_from_docx(book_name, docx_path):
    # Helpers: italics-path check for quote enclosure (compare-only, no text mutation)
    import re as _re_local
    _SPACE_LIKE = _re_local.compile(r'[\u0020\u00A0\u2009\u200A\u200B\u202F\u205F\u3000]+')
    _OPEN_QS  = {'“', '"'}
    _CLOSE_QS = {'”', '"'}
    def _prev_non_space(_s: str, _idx: int) -> str:
        i = _idx - 1
        while i >= 0 and _SPACE_LIKE.match(_s[i]):
            i -= 1
        return _s[i] if i >= 0 else ''
    def _next_non_space(_s: str, _idx: int) -> str:
        n = len(_s); i = _idx
        while i < n and _SPACE_LIKE.match(_s[i]):
            i += 1
        return _s[i] if i < n else ''
    def _is_enclosed_by_quotes(_para_text: str, _start: int, _end: int, _seg_text: str) -> bool:
        # Pattern A: quotes outside the italic span: … “ [italic] ” …
        _left  = _prev_non_space(_para_text, _start)
        _right = _next_non_space(_para_text, _end)
        if _left in _OPEN_QS and _right in _CLOSE_QS:
            return True
        # Pattern B: quotes inside the italic span text (rare)
        st = _seg_text
        if st and st[0] in _OPEN_QS:
            j = 1
            while j < len(st) and _SPACE_LIKE.match(st[j]):
                j += 1
            st = st[:1] + st[j:]
        if st and st[-1] in _CLOSE_QS:
            k = len(st) - 2
            while k >= 0 and _SPACE_LIKE.match(st[k]):
                k -= 1
            st = st[:k+1] + st[-1:]
        return (len(st) >= 2 and st[0] in _OPEN_QS and st[-1] in _CLOSE_QS)
    doc = docx.Document(docx_path)
    quote_pattern = re.compile(r'(?:^|\s)(["“].+?["”])(?=$|[\s\.\,\;\:\!\?\)\]\}])')
    dialogue_list = []
    line_number = 1
    for para in doc.paragraphs:
        text = para.text.strip()
    
        # Build ordered segments: closing-only -> paired -> opening-only
        matches = list(quote_pattern.finditer(text))  # paired matches
        covered = []

        def _overlaps(a, b):
            # a, b are (start, end) half-open intervals
            return not (a[1] <= b[0] or b[1] <= a[0])

        def _is_covered(span):
            return any(_overlaps(span, c) for c in covered)

        OPEN = {'“', '"'}
        CLOSE = {'”', '"'}

        ordered = []

        # 1) Closing-only: first closing before any opening in *uncovered* text
        first_close = -1
        first_open = -1
        for i, ch in enumerate(text):
            if _is_covered((i, i + 1)):
                continue
            if ch in CLOSE and first_close == -1:
                first_close = i
            if ch in OPEN and first_open == -1:
                first_open = i
            if first_close != -1 and (first_open == -1 or first_close < first_open):
                seg_span = (0, first_close + 1)
                seg = text[seg_span[0]:seg_span[1]].strip()
                if seg:
                    ordered.append((seg_span, seg))
                    covered.append(seg_span)
                break  # only the first closing-only segment per paragraph

        # 2) Paired quotes: use existing regex; skip spans already covered
        for m in matches:
            seg_span = (m.start(1), m.end(1))
            if not _is_covered(seg_span):
                seg = m.group(1).strip()
                if seg:
                    ordered.append((seg_span, seg))
                    covered.append(seg_span)

        # 3) Opening-only: last opening with no closing after -> opening..end
        last_open = -1
        for i, ch in enumerate(text):
            if ch in OPEN and not _is_covered((i, i + 1)):
                last_open = i

        if last_open != -1:
            # any *uncovered* closing after last_open?
            has_close_after = False
            j = last_open + 1
            while j < len(text):
                if not _is_covered((j, j + 1)) and text[j] in CLOSE:
                    has_close_after = True
                    break
                j += 1
            if not has_close_after:
                seg_span = (last_open, len(text))
                if not _is_covered(seg_span):
                    seg = text[seg_span[0]:seg_span[1]].strip()
                    if seg:
                        ordered.append((seg_span, seg))
                        covered.append(seg_span)

        # Merge quotes (with spans) and italics (with spans), then sort by reading order
        items = []  # list of ((start, end), text)

        # quotes collected earlier as (span, text)
        for span, seg in ordered:
            items.append((span, seg))
        
        # italics: use content-only spans to avoid duplicating fully italicised quotes
        quote_spans = [span for span, _ in ordered]
        def _inside_any(inner_span, outer_spans):
            s, e = inner_span
            return any(os <= s and e <= oe for (os, oe) in outer_spans)
        
        for i_span, i_text in extract_italic_spans(para):
            if _inside_any(i_span, quote_spans):
                continue
            items.append((i_span, i_text))
        # sort: start asc, then longer span first (desc)
        
        # Emit italics only if their *content* span lies outside all quote spans
        
        # Helper: trim leading/trailing quote characters from a span for containment checks
        def _trim_quotes_for_span(span, src_text):
            s, e = span
            QUOTES = {'“','”','"','\'','‘','’'}
            while s < e and src_text[s] in QUOTES:
                s += 1
            while e > s and src_text[e-1] in QUOTES:
                e -= 1
            return (s, e)
for im in re.finditer(r"<i>(.*?)</i>", text, flags=re.DOTALL):
            i_span = (im.start(1), im.end(1))  # content-only span (excludes tags)
            i_text = im.group(1)
            t_span = _trim_quotes_for_span(i_span, text)
            # Skip italics if either raw span or trimmed span lies inside any quote
            if _inside_any(i_span, quote_spans) or _inside_any(t_span, quote_spans):
                continue
            items.append((i_span, i_text))
items.sort(key=lambda it: (it[0][0], -(it[0][1] - it[0][0])))

        for _, seg in items:
            dialogue_list.append(f"{line_number}. Unknown: {seg}")
            line_number += 1
    output_path = f"{st.session_state.userkey}-{book_name}-quotes.txt"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(dialogue_list))
    return dialogue_list

# ---------------------------
# DOCX-to-HTML & Marking Functions
# ---------------------------
def prepend_marker_to_paragraph(paragraph, marker_text):
    p = paragraph._p
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = marker_text + " "
    r.append(t)
    p.insert(0, r)

def create_marker_docx(original_docx, marker_docx):
    doc = docx.Document(original_docx)
    for idx, para in enumerate(doc.paragraphs):
        marker = f"[[[P{idx}]]]"
        prepend_marker_to_paragraph(para, marker)
    doc.save(marker_docx)

def convert_docx_to_html_mammoth(docx_file):
    with open(docx_file, "rb") as f:
        result = mammoth.convert_to_html(f)
        return result.value

def get_manual_indentation(docx_file):
    doc = docx.Document(docx_file)
    indented_paras = {}
    for idx, para in enumerate(doc.paragraphs):
        left = para.paragraph_format.left_indent
        right = para.paragraph_format.right_indent
        if (left is not None and left.pt > 0) or (right is not None and right.pt > 0):
            indented_paras[idx] = (left, right)
    return indented_paras

def convert_length_to_px(length):
    return length.pt * 1.33 if length is not None else 0

# ---------------------------
# Dialogue Highlighting Functions
# ---------------------------
def highlight_across_nodes(parent, quote, highlight_style, soup):
    full_text = parent.get_text()
    full_text_lower = match_normalize(full_text).lower()
    quote_lower = match_normalize(quote).lower()
    idx = full_text_lower.find(quote_lower)
    if idx == -1:
        return False
    end_idx = idx + len(quote)
    running_index = 0
    for descendant in list(parent.descendants):
        if isinstance(descendant, NavigableString):
            text = str(descendant)
            text_length = len(text)
            node_start = running_index
            node_end = running_index + text_length
            if node_end > idx and node_start < end_idx:
                overlap_start = max(idx, node_start)
                overlap_end = min(end_idx, node_end)
                rel_start = overlap_start - node_start
                rel_end = overlap_end - node_start
                before = text[:rel_start]
                match_text = text[rel_start:rel_end]
                after = text[rel_end:]
                new_nodes = []
                if before:
                    new_nodes.append(NavigableString(before))
                span_tag = soup.new_tag("span", attrs={"class": "highlight", "style": highlight_style})
                span_tag.string = match_text
                new_nodes.append(span_tag)
                if after:
                    new_nodes.append(NavigableString(after))
                descendant.replace_with(*new_nodes)
            running_index += text_length
    return True

def highlight_quote_in_parent(parent, quote, highlight_style, soup):
    stripped_quote = quote.strip('“”"')
    stripped_quote_lower = match_normalize(stripped_quote).lower()
    for child in parent.contents:
        if isinstance(child, NavigableString):
            child_text = str(child)
            child_text_lower = match_normalize(child_text).lower()
            idx = child_text_lower.find(stripped_quote_lower)
            if idx != -1:
                before = child_text[:idx]
                match_text = child_text[idx: idx + len(stripped_quote)]
                after = child_text[idx + len(stripped_quote):]
                new_nodes = []
                if before:
                    new_nodes.append(NavigableString(before))
                span_tag = soup.new_tag("span", attrs={"class": "highlight", "style": highlight_style})
                span_tag.string = match_text
                new_nodes.append(span_tag)
                if after:
                    new_nodes.append(NavigableString(after))
                child.replace_with(*new_nodes)
                return True
        elif hasattr(child, 'contents'):
            if highlight_quote_in_parent(child, quote, highlight_style, soup):
                return True
    return highlight_across_nodes(parent, stripped_quote, highlight_style, soup)

def build_candidate_info(soup):
    candidates = soup.find_all(['p', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
    candidate_info = []
    global_offset = 0
    for candidate in candidates:
        text = candidate.get_text()
        length = len(text)
        candidate_info.append((candidate, global_offset, global_offset + length, text))
        global_offset += length
    return candidate_info

def highlight_in_candidate(candidate, quote, highlight_style, soup, start_offset=0):
    full_text = candidate.get_text()
    full_text_lower = match_normalize(full_text).lower()
    quote_lower = match_normalize(quote).lower()
    pos = full_text_lower.find(quote_lower, start_offset)
    if pos == -1:
        return None
    match_end = pos + len(quote)
    running_index = 0
    for descendant in list(candidate.descendants):
        if isinstance(descendant, NavigableString):
            text = str(descendant)
            text_length = len(text)
            node_start = running_index
            node_end = running_index + text_length
            if node_end > pos and node_start < match_end:
                overlap_start = max(pos, node_start)
                overlap_end = min(match_end, node_end)
                rel_start = overlap_start - node_start
                rel_end = overlap_end - node_start
                before = text[:rel_start]
                match_text = text[rel_start:rel_end]
                after = text[rel_end:]
                new_nodes = []
                if before:
                    new_nodes.append(NavigableString(before))
                span_tag = soup.new_tag("span", attrs={"class": "highlight", "style": highlight_style})
                span_tag.string = match_text
                new_nodes.append(span_tag)
                if after:
                    new_nodes.append(NavigableString(after))
                descendant.replace_with(*new_nodes)
            running_index += text_length
    return match_end

def highlight_dialogue_in_html(html, quotes_list, speaker_colors):
    soup = BeautifulSoup(html, "html.parser")
    candidate_info = build_candidate_info(soup)
    unmatched_quotes = []
    last_global_offset = 0
    for quote_data in quotes_list:
        expected_quote = quote_data['quote'].strip('“”"')
        expected_quote_lower = match_normalize(expected_quote).lower()
        speaker = quote_data['speaker']
        norm_speaker = normalize_speaker_name(speaker)
        color_choice = st.session_state.speaker_colors.get(norm_speaker, "none")
        if norm_speaker == "unknown":
            color_choice = "none"
        rgba = COLOR_PALETTE.get(color_choice, COLOR_PALETTE["none"])
        if color_choice == "none":
            highlight_style = f"color: rgb({rgba[0]}, {rgba[1]}, {rgba[2]}); background-color: transparent;"
        else:
            highlight_style = f"color: {rgba[4]}; background-color: rgba({rgba[0]}, {rgba[1]}, {rgba[2]}, {rgba[3]});"
        matched = False
        for candidate, start, end, text in candidate_info:
            if end < last_global_offset:
                continue
            local_start = last_global_offset - start if last_global_offset > start else 0
            candidate_text_norm = match_normalize(text).lower()
            pos = candidate_text_norm.find(expected_quote_lower, local_start)
            if pos != -1:
                match_end_local = highlight_in_candidate(candidate, quote_data['quote'], highlight_style, soup, local_start)
                if match_end_local is not None:
                    last_global_offset = start + match_end_local
                    matched = True
                    break
        if not matched:
            for candidate, start, end, text in candidate_info:
                candidate_text_norm = match_normalize(text).lower()
                pos = candidate_text_norm.find(expected_quote_lower)
                if pos != -1:
                    match_end_local = highlight_in_candidate(candidate, quote_data['quote'], highlight_style, soup, 0)
                    if match_end_local is not None:
                        if start + match_end_local > last_global_offset:
                            last_global_offset = start + match_end_local
                        matched = True
                        break
        if not matched:
            unmatched_quotes.append(f"{quote_data['speaker']}: \"{quote_data['quote']}\" [Index: {quote_data['index']}]")
    if unmatched_quotes:
        unmatched_quotes_filename = get_unmatched_quotes_filename()
        with open(unmatched_quotes_filename, "w", encoding="utf-8") as f:
            f.write("\n".join(unmatched_quotes))
        st.write(f"⚠️ Unmatched quotes saved to '[userkey]-unmatched_quotes.txt' ({len(unmatched_quotes)} entries)")
    return str(soup)

def apply_manual_indentation_with_markers(original_docx, html):
    indented_paras = get_manual_indentation(original_docx)
    soup = BeautifulSoup(html, "html.parser")
    marker_regex = re.compile(r"\[\[\[P(\d+)\]\]\]")
    candidate_tags = ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']
    for tag in soup.find_all(candidate_tags):
        text = tag.get_text()
        match = marker_regex.search(text)
        if match:
            para_index = int(match.group(1))
            for text_node in tag.find_all(string=True):
                if marker_regex.search(text_node):
                    new_text = marker_regex.sub("", text_node)
                    text_node.replace_with(new_text)
            if para_index in indented_paras:
                left, right = indented_paras[para_index]
                left_px = convert_length_to_px(left)
                right_px = convert_length_to_px(right)
                style_str = f"margin-left: {left_px}px; margin-right: {right_px}px;"
                if tag.has_attr("style"):
                    tag["style"] += " " + style_str
                else:
                    tag["style"] = style_str
    return str(soup)

# ---------------------------
# Summary & Ranking Functions
# ---------------------------
def generate_summary_html(quotes_list, speakers, speaker_colors):
    counts = Counter(quote["speaker"] for quote in quotes_list)
    total_lines = sum(counts.values())
    summary_order = []
    if "Unknown" in counts:
        summary_order.append("Unknown")
    for sp in speakers:
        if sp != "Unknown" and sp not in summary_order:
            summary_order.append(sp)
    for sp in counts:
        if sp not in summary_order:
            summary_order.append(sp)
    lines = []
    lines.append('<div id="character-summary" style="border: 1px solid #ccc; padding: 10px; margin-bottom: 20px;">')
    lines.append('<h2 style="margin: 0 0 5px 0;">Character Summary</h2>')
    for sp in summary_order:
        if normalize_speaker_name(sp) in ("error", "do not read"):
            continue
        count = counts.get(sp, 0)
        percentage = round((count / total_lines) * 100) if total_lines > 0 else 0
        color_key = speaker_colors.get(normalize_speaker_name(sp), "none")
        if sp.lower() == "unknown":
            color_key = "none"
        rgba = COLOR_PALETTE.get(color_key, COLOR_PALETTE["none"])
        if color_key == "none":
            style = f"color: rgb({rgba[0]}, {rgba[1]}, {rgba[2]}); background-color: transparent;"
        else:
            style = f"color: {rgba[4]}; background-color: rgba({rgba[0]}, {rgba[1]}, {rgba[2]}, {rgba[3]});"
        lines.append(f'<p style="margin: 0; line-height: 1.2; padding: 8px 0;"><span class="highlight" style="{style}">{sp}</span> - {count} lines - {percentage}%</p>')
    lines.append('</div>')
    return "\n".join(lines)

def generate_ranking_html(quotes_list, speaker_colors):
    counts = Counter(quote["speaker"] for quote in quotes_list)
    total_lines = sum(counts.values())
    filtered = [(sp, count) for sp, count in counts.items() if sp.lower() not in ("unknown", "do not read", "error") and count > 1]
    filtered.sort(key=lambda x: x[1], reverse=True)
    lines = []
    lines.append('<div id="speaker-ranking" style="margin-top: 20px;">')
    lines.append('<h2 style="margin: 0 0 5px 0;">Speaker Ranking</h2>')
    for sp, count in filtered:
        percentage = round((count / total_lines) * 100) if total_lines > 0 else 0
        color_key = speaker_colors.get(normalize_speaker_name(sp), "none")
        rgba = COLOR_PALETTE.get(color_key, COLOR_PALETTE["none"])
        if color_key == "none":
            style = f"color: rgb({rgba[0]}, {rgba[1]}, {rgba[2]}); background-color: transparent;"
        else:
            style = f"color: {rgba[4]}; background-color: rgba({rgba[0]}, {rgba[1]}, {rgba[2]}, {rgba[3]});"
        lines.append(f'<p style="margin: 0; line-height: 1.2; padding: 8px 0;"><span class="highlight" style="{style}">{sp}</span> - {count} lines - {percentage}%</p>')
    lines.append('</div>')
    return "\n".join(lines)

def generate_first_lines_html(quotes_list, speakers):
    # Map: speaker (canonical) -> first qualifying quote
    first_lines = {}
    for quote in quotes_list:
        speaker = quote["speaker"]
        norm = normalize_speaker_name(speaker)
        # Only add the first qualifying line (3+ words, else first)
        if norm not in first_lines:
            words = quote["quote"].strip().split()
            if len(words) >= 3 or all(q["speaker"] != speaker for q in quotes_list if q != quote):
                first_lines[norm] = quote["quote"].strip()
            else:
                # Tentatively store, may be replaced by a later 3+ word line
                first_lines[norm] = quote["quote"].strip()
        else:
            if len(first_lines[norm].split()) < 3 and len(quote["quote"].strip().split()) >= 3:
                first_lines[norm] = quote["quote"].strip()

    lines = []
    lines.append('<div id="first-lines-summary" style="border: 1px solid #ccc; padding: 10px; margin-bottom: 20px;">')
    lines.append('<h2 style="margin: 0 0 5px 0;">First Substantial Lines</h2>')
    for sp in speakers:
        norm = normalize_speaker_name(sp)
        if norm in ("do not read", "error", "unknown"):
            continue
        if norm in first_lines:
            line = first_lines[norm]
            lines.append(f'<p style="margin: 0; line-height: 1.2; padding: 8px 0;"><span class="highlight">{sp}</span>: <span style="font-style: italic;">{line}</span></p>')
    lines.append('</div>')
    return "\n".join(lines)
  
# ---------------------------
# Canonical Speaker & Quote Functions
# ---------------------------
def get_canonical_speakers(quotes_file):
    speakers = []
    pattern = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
    with open(quotes_file, "r", encoding="utf-8") as f:
        for line in f:
            match = pattern.match(line.strip())
            if match:
                speaker_raw = match.group(1).strip()
                speakers.append(smart_title(str(speaker_raw)))  # Ensure it's a string
    seen = set()
    canonical_speakers = []
    for s in speakers:
        norm = normalize_speaker_name(str(s))
        if norm not in seen:
            seen.add(norm)
            canonical_speakers.append(s)
    canonical_map = {normalize_speaker_name(str(s)): s for s in canonical_speakers}
    return canonical_speakers, canonical_map

def load_quotes(quotes_file, canonical_map):
    quotes_list = []
    pattern = re.compile(r"^\s*([0-9]+(?:[a-zA-Z]+)?)\.\s+([^:]+):\s*(?:[“\"])?(.+?)(?:[”\"])?\s*$")
    with open(quotes_file, "r", encoding="utf-8") as f:
        for line in f:
            match = pattern.match(line.strip())
            if match:
                index, speaker_raw, quote = match.groups()
                effective = smart_title(speaker_raw)
                norm = normalize_speaker_name(effective)
                canonical = canonical_map.get(norm, effective)
                quotes_list.append({
                    "index": index,
                    "speaker": canonical,
                    "quote": quote.strip()
                })
    return quotes_list

def load_existing_colors():
  if os.path.exists(get_saved_colors_file()):
    with open(get_saved_colors_file(), "r", encoding="utf-8") as f:
        loaded_colors = json.load(f)
    normalized_loaded = {normalize_speaker_name(k): v for k, v in loaded_colors.items()}
    st.session_state.speaker_colors = normalized_loaded
    st.session_state.existing_speaker_colors = normalized_loaded

def save_speaker_colors(speaker_colors):
    with open(get_saved_colors_file(), "w", encoding="utf-8") as f:
        json.dump(speaker_colors, f, indent=4, ensure_ascii=False)

# ---------------------------
# Restart Helper Function
# ---------------------------
def restart_app():
    st.session_state.clear()
    st.rerun()

# ---------------------------
# Streamlit Multi-Step UI
# ---------------------------

# ========= STEP 1: Upload & Initialize =========
if st.session_state.step == 1:
    st.markdown("<h4>DOCX to HTML Converter with Dialogue Highlighting</h4>", unsafe_allow_html=True)
    st.write("Upload your DOCX and quotes text files. Optionally, upload an existing speaker_colors.json file.")
    st.write("Alternatively, upload **just a DOCX** to create a quotes text file.")
    
    if "docx_bytes" in st.session_state:
        st.success("DOCX already uploaded and processed.")
        if st.session_state.get("docx_only", False):
            if st.session_state.get("quotes_lines") is not None:
                quotes_txt = "\n".join(st.session_state.quotes_lines)
                st.download_button("Download Extracted Quotes TXT", quotes_txt.encode("utf-8"),
                                   file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt", mime="text/plain")
                if st.button("Restart", key="restart_docx"):
                    restart_app()
                if st.button("Continue", key="continue_docx"):
                    st.session_state.docx_only = False
                    st.session_state.unknown_index = 0
                    st.session_state.console_log = []
                    st.session_state.step = 2
                    # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
                    try:
                        if st.session_state.get("quotes_lines"):
                            pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                            counts_cap10 = {}
                            flagged = set()
                            for _line in st.session_state.quotes_lines:
                                m = pattern_speaker.match(_line.strip())
                                if not m:
                                    continue
                                speaker_raw = m.group(1).strip()
                                effective = smart_title(speaker_raw)
                                norm = normalize_speaker_name(effective)
                                if norm in flagged:
                                    continue
                                c = counts_cap10.get(norm, 0)
                                if c < 10:
                                    c += 1
                                    counts_cap10[norm] = c
                                    if c >= 10:
                                        flagged.add(norm)
                            st.session_state.speaker_counts = counts_cap10
                            st.session_state.flagged_names = flagged
                        else:
                            if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                                st.session_state.speaker_counts = {}
                            if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                                st.session_state.flagged_names = set()
                    except Exception:
                        pass
                    auto_save()
                    st.rerun()
            else:
                dialogue_list = extract_dialogue_from_docx(st.session_state.book_name, st.session_state.docx_path)
                st.session_state.quotes_lines = [line + "\n" for line in dialogue_list]
                st.session_state.docx_only = True
                st.success("Quotes extracted from DOCX.")
                quotes_txt = "\n".join(dialogue_list)
                st.download_button("Download Extracted Quotes TXT", quotes_txt.encode("utf-8"),
                                   file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt", mime="text/plain")
                if st.button("Restart", key="restart_docx"):
                    restart_app()
                if st.button("Continue", key="continue_docx"):
                    st.session_state.docx_only = False
                    st.session_state.unknown_index = 0
                    st.session_state.console_log = []
                    st.session_state.step = 2
                    # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
                    try:
                        if st.session_state.get("quotes_lines"):
                            pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                            counts_cap10 = {}
                            flagged = set()
                            for _line in st.session_state.quotes_lines:
                                m = pattern_speaker.match(_line.strip())
                                if not m:
                                    continue
                                speaker_raw = m.group(1).strip()
                                effective = smart_title(speaker_raw)
                                norm = normalize_speaker_name(effective)
                                if norm in flagged:
                                    continue
                                c = counts_cap10.get(norm, 0)
                                if c < 10:
                                    c += 1
                                    counts_cap10[norm] = c
                                    if c >= 10:
                                        flagged.add(norm)
                            st.session_state.speaker_counts = counts_cap10
                            st.session_state.flagged_names = flagged
                        else:
                            if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                                st.session_state.speaker_counts = {}
                            if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                                st.session_state.flagged_names = set()
                    except Exception:
                        pass
                    auto_save()
                    st.rerun()
        else:
            pass
    else:
        docx_file = st.file_uploader("Upload DOCX File", type=["docx"])
        quotes_file = st.file_uploader("Upload Quotes TXT File (optional)", type=["txt"])
        speaker_colors_file = st.file_uploader("Upload Speaker Colors JSON (optional)", type=["json"])
        
        if st.button("Start Processing"):
            if docx_file is None:
                st.error("Please upload a DOCX file.")
            else:
                st.session_state.book_name = os.path.splitext(docx_file.name)[0]
                st.session_state.docx_bytes = docx_file.getvalue()
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                    tmp_docx.write(st.session_state.docx_bytes)
                    st.session_state.docx_path = tmp_docx.name
                if speaker_colors_file is not None:
                    raw = json.load(speaker_colors_file)
                    st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in raw.items()}
                    save_speaker_colors(st.session_state.existing_speaker_colors)
                    st.session_state.speaker_colors = st.session_state.existing_speaker_colors.copy()
                else:
                    st.session_state.existing_speaker_colors = {}
                    st.session_state.speaker_colors = {}
                if quotes_file is not None:
                    quotes_text = quotes_file.read().decode("utf-8")
                    st.session_state.quotes_lines = quotes_text.splitlines(keepends=True)
                    st.session_state.docx_only = False
                    # Persist uploaded quotes to a consistent filename and ensure JSON cache for previews
                    try:
                        # Save uploaded quotes to working directory with userkey-bookname naming
                        if 'userkey' in st.session_state and st.session_state.get('book_name'):
                            quotes_filename = f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt"
                        else:
                            # Fallback name if userkey/book_name not set for any reason
                            quotes_filename = "uploaded-quotes.txt"
                        with open(quotes_filename, "w", encoding="utf-8") as _qf:
                            _qf.write(quotes_text)
                        # Ensure the paragraph JSON exists (no DOCX fallback in Step 2)
                        write_paragraph_json_for_session()
                    except Exception as _e:
                        # Do not fail hard; Step 2 will show 'JSON cache not found yet' if this fails
                        pass
                else:
                    st.session_state.quotes_lines = None
                    st.session_state.docx_only = True
                
                    # Create/overwrite the paragraph JSON once here for docx-only case
                    write_paragraph_json_for_session()
                st.session_state.unknown_index = 0
                st.session_state.console_log = []
                if st.session_state.docx_only:
                    st.session_state.step = 1
                else:
                    st.session_state.step = 2
                    # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
                    try:
                        if st.session_state.get("quotes_lines"):
                            pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                            counts_cap10 = {}
                            flagged = set()
                            for _line in st.session_state.quotes_lines:
                                m = pattern_speaker.match(_line.strip())
                                if not m:
                                    continue
                                speaker_raw = m.group(1).strip()
                                effective = smart_title(speaker_raw)
                                norm = normalize_speaker_name(effective)
                                if norm in flagged:
                                    continue
                                c = counts_cap10.get(norm, 0)
                                if c < 10:
                                    c += 1
                                    counts_cap10[norm] = c
                                    if c >= 10:
                                        flagged.add(norm)
                            st.session_state.speaker_counts = counts_cap10
                            st.session_state.flagged_names = flagged
                        else:
                            if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                                st.session_state.speaker_counts = {}
                            if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                                st.session_state.flagged_names = set()
                    except Exception:
                        pass
                auto_save()
                st.rerun()

# ========= STEP 2: Unknown Speaker Processing =========
elif st.session_state.step == 2:
    st.markdown("<h4>Step 2: Process Unknown Speakers</h4>", unsafe_allow_html=True)
    st.write("For each quote with speaker 'Unknown', type a replacement (or type 'skip', 'exit', or 'undo').")
    
    def get_next_unknown_line():
        quotes = st.session_state.get("quotes_lines")
        if quotes is None:
            return None, None, None
        pattern = re.compile(r"^(\s*\d+(?:[a-zA-Z]+)?\.\s+)([^:]+)(:.*)$")
        for i in range(st.session_state.unknown_index, len(quotes)):
            line = quotes[i]
            m = pattern.match(line)
            if m:
                prefix, speaker_raw, remainder = m.groups()
                if speaker_raw.strip() == "Unknown":
                    return i, prefix, remainder
        return None, None, None

    index, prefix, remainder = get_next_unknown_line()
    if index is None:
        st.write("No more unknown speakers found.")
        if st.button("Proceed to Color Assignment"):
            st.session_state.step = 3
            auto_save()
            st.rerun()
    else:
        dialogue = remainder.lstrip(": ").rstrip("\n")
        st.markdown("<hr style='margin: 2px 0;'>", unsafe_allow_html=True)
        # Using global JSON-only context resolver
        
        # Compute occurrence target from previous two lines in quotes (quoted-segment aware)
        occurrence_target = 1
        try:
            qlines = st.session_state.get("quotes_lines") or []
            patt = re.compile(r"^(\s*\d+(?:[a-zA-Z]+)?\.\s+)([^:]+)(:.*)$")
            def remainder_for(i):
                if i is None or i < 0 or i >= len(qlines):
                    return None
                mm = patt.match(qlines[i])
                if not mm:
                    return None
                return mm.group(3).lstrip(": ").rstrip("\n")
            def first_quoted_segment(s: str) -> str:
                if s is None:
                    return ""
                m1 = re.search(r'[“"]([^”"]+)[”"]', s)
                if not m1:
                    m1 = re.search(r"[‘']([^’']+)[’']", s)
                return m1.group(1) if m1 else s
            curr_seg = first_quoted_segment(dialogue)
            curr_norm = normalize_text(curr_seg).lower()
            prev1 = remainder_for(index-1)
            prev2 = remainder_for(index-2)
            prev1_norm = normalize_text(first_quoted_segment(prev1)).lower() if prev1 else ""
            prev2_norm = normalize_text(first_quoted_segment(prev2)).lower() if prev2 else ""
            def _norm_contains(a: str, b: str) -> bool:
                try:
                    if not a or not b:
                        return False
                    # One-way containment: treat as repeat only if CURRENT (a) is within PREVIOUS (b)
                    return a in b
                except Exception:
                    return False
            
            count_prev_same = int(_norm_contains(curr_norm, prev1_norm)) + int(_norm_contains(curr_norm, prev2_norm))            
            occurrence_target = 1 + count_prev_same
#            st.session_state._dbg_occurrence_target = occurrence_target
#            st.session_state._dbg_prev1_norm = prev1_norm
#            st.session_state._dbg_prev2_norm = prev2_norm
#            st.session_state._dbg_curr_norm = curr_norm
        except Exception:
            pass
        context = get_context_for_dialogue_json_only(dialogue, occurrence_target=occurrence_target)
        if context:
            # Remember the currently displayed previous paragraph for potential trimming upon match
            try:
                st.session_state.context_previous_candidate = context.get("previous")
            except Exception:
                st.session_state.context_previous_candidate = None
            if "previous" in context:
                st.markdown(neutralize_markdown_in_html(context["previous"]), unsafe_allow_html=True)
            st.markdown(neutralize_markdown_in_html(context["current"]), unsafe_allow_html=True)
#            try:
#                st.text('DEBUG CMP: curr=' + str(st.session_state.get('_dbg_curr_norm'))
#                         + ' | prev1=' + str(st.session_state.get('_dbg_prev1_norm'))
#                         + ' | prev2=' + str(st.session_state.get('_dbg_prev2_norm')))
#            except Exception:
#                pass

            if "next" in context:
                st.markdown(neutralize_markdown_in_html(context["next"]), unsafe_allow_html=True)
        else:
            st.write("No context found in cached JSON for this quote.")
        st.markdown("<hr style='margin: 2px 0;'>", unsafe_allow_html=True)
        st.write(f"**Dialogue (Line {index+1}):** {dialogue}")
        
        def process_unknown_input(new_speaker: str):
            new_speaker = new_speaker.strip()
            if not new_speaker:
                st.session_state.console_log.insert(0, "Empty input ignored. Enter a name, or use 'skip' / 'exit'.")
                return
            if new_speaker.lower() == "exit":
                st.session_state.console_log.insert(0, "Exiting unknown speaker processing.")
                st.session_state.step = 3
            elif new_speaker.lower() == "skip":
                st.session_state.console_log.insert(0, f"Skipped line {index+1}.")
                st.session_state.unknown_index = index + 1
            elif new_speaker.lower() == "undo":
                if "last_update" in st.session_state:
                    last_index = st.session_state.last_update[0]
                    pattern = re.compile(r"^(\s*\d+(?:[a-zA-Z]+)?\.\s+)([^:]*)(:.*)$")
                    m = pattern.match(st.session_state.quotes_lines[last_index])
                    if m:
                        prefix_u, _, remainder_u = m.groups()
                        st.session_state.quotes_lines[last_index] = prefix_u + "Unknown" + remainder_u
                        st.session_state.unknown_index = last_index
                        st.session_state.console_log.insert(0, f"Reverted line {last_index+1} to Unknown.")
                    del st.session_state.last_update
                else:
                    st.session_state.console_log.insert(0, "Nothing to undo.")
            else:
                st.session_state.last_update = (index, st.session_state.quotes_lines[index])
                # On confirmed match only (not skip/exit/undo), trim paragraph cache before the "previous" that was displayed.
                try:
                    prev_for_trim = st.session_state.get("context_previous_candidate")
                except Exception:
                    prev_for_trim = None
                try:
                    _trim_ok = trim_paragraph_cache_before_previous(prev_for_trim)
                    if _trim_ok:
                        st.session_state.console_log.insert(0, "Trimmed paragraph cache before previous context.")
                except Exception:
                    pass

                updated_speaker = smart_title(new_speaker)

                # Increment count for unflagged speakers and flag at 10
                try:
                    norm = normalize_speaker_name(updated_speaker)
                    if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                        st.session_state.speaker_counts = {}
                    if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                        st.session_state.flagged_names = set()
                    if norm not in st.session_state.flagged_names:
                        new_cnt = st.session_state.speaker_counts.get(norm, 0) + 1
                        if new_cnt >= 10:
                            new_cnt = 10
                            st.session_state.flagged_names.add(norm)
                        st.session_state.speaker_counts[norm] = new_cnt
                except Exception as _e:
                    pass
                new_line = prefix + updated_speaker + remainder
                if not new_line.endswith("\n"):
                    new_line += "\n"
                st.session_state.quotes_lines[index] = new_line
                st.session_state.console_log.insert(0, f"Updated line {index+1} with speaker: {updated_speaker}")
                st.session_state.unknown_index = index + 1
            auto_save()
            st.rerun()
        
        # --- New: one‑submit‑per‑name form -------------------------------

        # Frequent speakers (flagged, alphabetical). Buttons act like typing + Enter.
        try:
            if "flagged_names" in st.session_state and st.session_state.flagged_names:
                flagged_sorted = sorted(st.session_state.flagged_names)
                flagged_sorted = [n for n in flagged_sorted if n.lower() != "unknown"]
                st.caption("Frequent speakers:")
                with st.container(horizontal=True):
                    cmap = st.session_state.get("canonical_map") or {}
                    for i, norm in enumerate(flagged_sorted):
                        display_name = cmap.get(norm, norm.title())
                        if st.button(display_name, key=f"flagged_{norm}"):
                            process_unknown_input(display_name)
        except Exception as _e:
            pass
        with st.form("unknown_form", clear_on_submit=True):
            new_name = st.text_input(
                "Enter speaker name (or 'skip'/'exit'/'undo'):",
                key="new_speaker_input",
                placeholder="Type name and press Enter",
            )
        
            # Evenly spaced horizontal buttons
            with st.container(horizontal=True):
                submitted    = st.form_submit_button("Submit")
                skip_clicked = st.form_submit_button("Skip")
                exit_clicked = st.form_submit_button("Exit")
                undo_clicked = st.form_submit_button("Undo (max 1)")
        if submitted:
            process_unknown_input(new_name)
        elif skip_clicked:
            process_unknown_input("skip")
        elif exit_clicked:
            process_unknown_input("exit")
        elif undo_clicked:
            process_unknown_input("undo")
          
        st.text_area("Console Log", "\n".join(st.session_state.console_log), height=150, label_visibility="collapsed")

# ========= STEP 3: Speaker Color Assignment =========
elif st.session_state.step == 3:
    st.markdown("<h4>Step 3: Speaker Color Assignment</h4>", unsafe_allow_html=True)
    st.write("Assign highlight colors for speakers that do not yet have an assigned color. You can also click 'Edit Speaker Colors' to review and change all assignments.")
    # Load the canonical speakers.
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w+", encoding="utf-8") as tmp_quotes:
        tmp_quotes.write("".join(st.session_state.quotes_lines))
        tmp_quotes_path = tmp_quotes.name
    canonical_speakers, canonical_map = get_canonical_speakers(tmp_quotes_path)
    st.session_state.canonical_map = canonical_map
    # Load existing colors (or default to empty dict)
    existing_colors = st.session_state.get("existing_speaker_colors") or load_existing_colors() or {}
    
    # Determine which speakers need a new assignment
    speakers_to_assign = [
        sp for sp in canonical_speakers 
        if sp.lower() not in ("unknown", "do not read") and (normalize_speaker_name(sp) not in existing_colors or existing_colors.get(normalize_speaker_name(sp), "none") == "none")
    ]
    
    if speakers_to_assign:
        st.write("Assign colors to the following speakers:")
        color_options = [color.title() for color in COLOR_PALETTE.keys() if color.lower() != "do not read"]
        updated_colors = {}
        for sp in speakers_to_assign:
            norm = normalize_speaker_name(sp)
            default_color = existing_colors.get(norm, "none")
            try:
                default_index = color_options.index(default_color.title())
            except ValueError:
                default_index   = color_options.index("None")
            selected = st.selectbox(sp, options=color_options, index=default_index, key="new_"+norm)
            updated_colors[norm] = selected.lower()
        # Merge updated colors with any already assigned values.
        for norm, col in updated_colors.items():
            existing_colors[norm] = col
        # Build final dictionary using normalized keys.
        final_colors = {}
        for sp in canonical_speakers:
            norm = normalize_speaker_name(sp)
            if sp.lower() == "unknown":
                final_colors[norm] = "none"
            elif sp.lower() == "do not read":
                final_colors[norm] = "do not read"
            else:
                final_colors[norm] = existing_colors.get(norm, "none")
        st.session_state.speaker_colors = final_colors
        st.session_state.existing_speaker_colors = existing_colors.copy()
        save_speaker_colors(final_colors)
        st.success("Speaker colors updated.")
    else:
        st.write("All speakers already have assigned colors.")
    if os.path.exists(get_saved_colors_file()):
        with open(get_saved_colors_file(), "r", encoding="utf-8") as f:
            loaded_colors = json.load(f)
        st.session_state.speaker_colors = loaded_colors
        st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in loaded_colors.items()}
    else:
        st.session_state.speaker_colors = {}
        st.session_state.existing_speaker_colors = {}
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Continue"):
            st.session_state.step = 4
            auto_save()
            st.rerun()
    with col2:
        if st.button("Edit Speaker Colors"):
            if os.path.exists(get_saved_colors_file()):
                with open(get_saved_colors_file(), "r", encoding="utf-8") as f:
                    loaded_colors = json.load(f)
                st.session_state.speaker_colors = loaded_colors
                st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in loaded_colors.items()}
            st.session_state.step = "edit_colors"
            auto_save()
            st.rerun()

# ========= EDIT COLORS: Full Speaker Color Assignment =========
elif st.session_state.step == "edit_colors":
    st.markdown("<h4>Edit Speaker Colors</h4>", unsafe_allow_html=True)
    st.write("Edit the assigned colors for all speakers (excluding 'Unknown'):")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w+", encoding="utf-8") as tmp_quotes:
        tmp_quotes.write("".join(st.session_state.quotes_lines))
        tmp_quotes_path = tmp_quotes.name
    canonical_speakers, canonical_map = get_canonical_speakers(tmp_quotes_path)
    st.session_state.canonical_map = canonical_map
    # Load current colors (or default to empty)
    existing_colors = st.session_state.get("speaker_colors") or load_existing_colors() or {}
    updated_colors = existing_colors.copy()
    color_options = [color.title() for color in COLOR_PALETTE.keys() if color.lower() != "do not read"]
    for sp in canonical_speakers:
        if sp.lower() in ("unknown", "do not read"):
            continue
        norm = normalize_speaker_name(sp)
        default_color = existing_colors.get(norm, "none")
        try:
            default_index = color_options.index(default_color.title())
        except ValueError:
            default_index = color_options.index("None")
        selected = st.selectbox(sp, options=color_options, index=default_index, key="edit_"+norm)
        updated_colors[norm] = selected.lower()
    st.session_state.speaker_colors = updated_colors
    st.session_state.existing_speaker_colors = updated_colors.copy()
    save_speaker_colors(updated_colors)
    st.success("Speaker colors updated.")
    if st.button("Continue"):
        st.session_state.step = 4
        auto_save()
        st.rerun()

# ========= STEP 4: Final HTML Generation =========
elif st.session_state.step == 4:
    if "speaker_colors" not in st.session_state or st.session_state.speaker_colors is None:
        st.session_state.speaker_colors = load_existing_colors() or {}
    st.markdown("<h4>Step 4: Final HTML Generation</h4>", unsafe_allow_html=True)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w+", encoding="utf-8") as tmp_quotes:
        tmp_quotes.write("".join(st.session_state.quotes_lines))
        quotes_file_path = tmp_quotes.name
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_marker:
        marker_docx_path = tmp_marker.name
    create_marker_docx(st.session_state.docx_path, marker_docx_path)
    html = convert_docx_to_html_mammoth(marker_docx_path)
    os.remove(marker_docx_path)
    quotes_list = load_quotes(quotes_file_path, st.session_state.canonical_map)
    highlighted_html = highlight_dialogue_in_html(html, quotes_list, st.session_state.speaker_colors)
    final_html_body = apply_manual_indentation_with_markers(st.session_state.docx_path, highlighted_html)
    summary_html = generate_summary_html(quotes_list, list(st.session_state.canonical_map.values()), st.session_state.speaker_colors)
    ranking_html = generate_ranking_html(quotes_list, st.session_state.speaker_colors)
    first_lines_html = generate_first_lines_html(quotes_list, list(st.session_state.canonical_map.values()))
    final_html_body = summary_html + "\n<br><br><br>\n" + ranking_html + "\n<br><br><br>\n" + first_lines_html + "\n" + final_html_body
    final_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{st.session_state.book_name}</title>
  <style>
    body {{
      font-family: Avenir, sans-serif;
      line-height: 2;
      max-width: 500px;
      margin: auto;
    }}
    span {{
      padding: 0;
    }}
    span.highlight {{
      background-color: var(--highlight-color, transparent);
      padding: 0.33em 0px;
      box-decoration-break: clone;
      -webkit-box-decoration-break: clone;
    }}
  </style>
</head>
<body>
{final_html_body}
</body>
</html>
"""
    final_html_path = os.path.join(tempfile.gettempdir(), f"{st.session_state.book_name}.html")
    with open(final_html_path, "w", encoding="utf-8") as f:
        f.write(final_html)
    st.success("Final HTML generated.")
    components.html(final_html, height=800, scrolling=True)
    with open(final_html_path, "rb") as f:
        html_bytes = f.read()
    st.download_button("Download HTML File", html_bytes,
                       file_name=f"{st.session_state.userkey}-{st.session_state.book_name}.html", mime="text/html")
    updated_colors = json.dumps(st.session_state.speaker_colors, indent=4, ensure_ascii=False).encode("utf-8")
    st.download_button("Download Updated Speaker Colors JSON", updated_colors,
                       file_name=f"{st.session_state.userkey}-speaker_colors.json", mime="application/json")
    updated_quotes = "".join(st.session_state.quotes_lines).encode("utf-8")
    st.download_button("Download Updated Quotes TXT", updated_quotes,
                       file_name=f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt", mime="text/plain")
    if os.path.exists(get_unmatched_quotes_filename()):
        with open(get_unmatched_quotes_filename(), "rb") as f:
            unmatched_bytes = f.read()
        st.download_button("Download Unmatched Quotes TXT", unmatched_bytes,
                           file_name=get_unmatched_quotes_filename(), mime="text/plain")
    if st.button("Return to Step 2"):
        if "book_name" in st.session_state:
            quotes_filename = f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt"
            if os.path.exists(quotes_filename):
                with open(quotes_filename, "r", encoding="utf-8") as f:
                    st.session_state.quotes_lines = f.read().splitlines(keepends=True)
        if os.path.exists(f"{st.session_state.userkey}-speaker_colors.json"):
            with open(f"{st.session_state.userkey}-speaker_colors.json", "r", encoding="utf-8") as f:
                colors = json.load(f)
            st.session_state.speaker_colors = colors
            st.session_state.existing_speaker_colors = {normalize_speaker_name(k): v for k, v in colors.items()}
        st.session_state.step = 2
        # Ensure frequent-speaker buttons are initialised from quotes before entering Step 2
        try:
            if st.session_state.get("quotes_lines"):
                pattern_speaker = re.compile(r"^\s*\d+(?:[a-zA-Z]+)?\.\s+([^:]+):")
                counts_cap10 = {}
                flagged = set()
                for _line in st.session_state.quotes_lines:
                    m = pattern_speaker.match(_line.strip())
                    if not m:
                        continue
                    speaker_raw = m.group(1).strip()
                    effective = smart_title(speaker_raw)
                    norm = normalize_speaker_name(effective)
                    if norm in flagged:
                        continue
                    c = counts_cap10.get(norm, 0)
                    if c < 10:
                        c += 1
                        counts_cap10[norm] = c
                        if c >= 10:
                            flagged.add(norm)
                st.session_state.speaker_counts = counts_cap10
                st.session_state.flagged_names = flagged
            else:
                if "speaker_counts" not in st.session_state or st.session_state.speaker_counts is None:
                    st.session_state.speaker_counts = {}
                if "flagged_names" not in st.session_state or st.session_state.flagged_names is None:
                    st.session_state.flagged_names = set()
        except Exception:
            pass
        auto_save()
        st.rerun() 
        # Add a Clear Cache button below "Return to Step 2"
    if st.button("Clear Cache for This User"):
        # Only delete files for this userkey
        files_to_remove = [
            get_progress_file(),
            get_saved_colors_file(),
            get_unmatched_quotes_filename(),
            f"{st.session_state.userkey}-{st.session_state.book_name}-quotes.txt",
            f"{st.session_state.userkey}-{st.session_state.book_name}.html",
            f"{st.session_state.userkey}-{st.session_state.book_name}.json"
        ]
        for path in files_to_remove:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception as e:
                st.warning(f"Could not remove {path}: {e}")

        # List all your app’s keys here to reset ONLY relevant state
        keys_to_clear = [
            "step", "userkey", "docx_bytes", "docx_path", "book_name",
            "quotes_lines", "speaker_colors", "existing_speaker_colors",
            "unknown_index", "console_log", "canonical_map", "last_update"
        ]
        for k in keys_to_clear:
            if k in st.session_state:
                del st.session_state[k]
        st.rerun()
